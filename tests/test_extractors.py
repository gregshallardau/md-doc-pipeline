"""Tests for document extraction logic."""

import pytest

from md_doc.extractors import extract_file


class TestPdfExtraction:
    def test_extract_pdf_returns_markdown(self, tmp_path):
        """PDF file should be converted to Markdown string."""
        pdf_path = tmp_path / "sample.pdf"

        # Create a minimal valid PDF
        pdf_bytes = bytes.fromhex(
            "255044462d312e340a312030206f626a0a3c3c202f54797065202f436174616c6f67202f5061676573203220302052203e3e0a656e646f626a0a322030206f626a0a3c3c202f54797065202f5061676573202f4b696473205b33203020525d202f436f756e742031203e3e0a656e646f626a0a332030206f626a0a3c3c202f54797065202f50616765202f506172656e74203220302052202f5265736f7572636573203c3c202f466f6e74203c3c202f4631203420302052203e3e203e3e202f4d65646961426f78205b30203020363132203739325d202f436f6e74656e7473203520302052203e3e0a656e646f626a0a342030206f626a0a3c3c202f54797065202f466f6e74202f53756274797065202f5479706531202f42617365466f6e74202f48656c766574696361203e3e0a656e646f626a0a352030206f626a0a3c3c202f4c656e677468203434203e3e0a73747265616d0a42540a2f46312031322054660a313030203730302054640a2854657374205044462920546a0a45540a656e6473747265616d0a656e646f626a0a787265660a3020360a303030303030303030302036353533352066200a30303030303030303039203030303030206e200a30303030303030303538203030303030206e200a30303030303030313135203030303030206e200a30303030303030323437203030303030206e200a30303030303030333334203030303030206e200a747261696c65720a3c3c202f53697a652036202f526f6f74203120302052203e3e0a7374617274787265660a3432380a2525454f460a"
        )
        pdf_path.write_bytes(pdf_bytes)

        result = extract_file(str(pdf_path))
        assert isinstance(result, str)
        assert len(result) > 0  # Non-empty markdown


class TestDocxExtraction:
    def test_extract_docx_returns_markdown(self, tmp_path):
        """DOCX file should be converted to Markdown string."""
        # Create a minimal DOCX using python-docx
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")

        docx_path = tmp_path / "sample.docx"
        doc.save(str(docx_path))

        result = extract_file(str(docx_path))
        assert isinstance(result, str)
        assert "Test Document" in result


class TestExtractFileTypeValidation:
    def test_extract_file_rejects_unsupported_type(self, tmp_path):
        """Unsupported file types should raise ValueError."""
        unsupported = tmp_path / "file.txt"
        unsupported.write_text("text file")

        with pytest.raises(ValueError, match="unsupported file type"):
            extract_file(str(unsupported))


class TestExtractFileNotFound:
    def test_extract_file_not_found(self, tmp_path):
        """Non-existent file should raise FileNotFoundError."""
        missing = tmp_path / "nonexistent.pdf"

        with pytest.raises(FileNotFoundError):
            extract_file(str(missing))
