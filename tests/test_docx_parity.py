"""Tests for PDF↔Word parity features in the DOCX/DOTX builder.

Covers the features that previously rendered only in PDF: body images,
mermaid diagrams, section heading bars, three-slot footers with page-number
fields, nested-list indentation, and graceful mermaid fallback.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from md_doc.builders.docx import build


@pytest.fixture()
def tmp_repo(tmp_path):
    (tmp_path / ".git").mkdir()
    return tmp_path


def _png(path: Path) -> None:
    from PIL import Image

    Image.new("RGB", (80, 40), "navy").save(path)


def _build(tmp_repo: Path, body: str, config: dict, fmt: str = "docx") -> Path:
    md = f"---\ntitle: T\n---\n\n{body}"
    doc_path = tmp_repo / "doc.md"
    doc_path.write_text(md, encoding="utf-8")
    out = tmp_repo / f"out.{fmt}"
    build(
        md,
        {"title": "T", "cover_page": False, **config},
        out,
        output_format=fmt,
        doc_path=doc_path,
        repo_root=tmp_repo,
    )
    return out


def _media_names(out: Path) -> list[str]:
    with zipfile.ZipFile(out) as z:
        return [n for n in z.namelist() if n.startswith("word/media/")]


def _part(out: Path, name: str) -> str:
    with zipfile.ZipFile(out) as z:
        return z.read(name).decode("utf-8")


def test_body_image_is_embedded(tmp_repo):
    _png(tmp_repo / "pic.png")
    out = _build(tmp_repo, "Text\n\n![logo](pic.png)\n", {})
    assert len(_media_names(out)) == 1


def test_unresolved_image_falls_back_to_alt_text(tmp_repo):
    out = _build(tmp_repo, "![my alt text](missing.png)\n", {})
    doc = Document(str(out))
    assert any("my alt text" in p.text for p in doc.paragraphs)
    assert _media_names(out) == []


def test_mermaid_diagram_is_rasterized(tmp_repo):
    pytest.importorskip("cairosvg")
    out = _build(tmp_repo, '```mermaid\npie\n"Done" : 100\n```\n', {})
    assert len(_media_names(out)) == 1
    # The diagram source must not leak as literal text.
    assert "language-mermaid" not in _part(out, "word/document.xml")


def test_mermaid_falls_back_to_code_when_no_rasterizer(tmp_repo, monkeypatch):
    import md_doc.builders._assets as assets

    monkeypatch.setattr(assets, "_svg_to_png", lambda *a, **k: None)
    out = _build(tmp_repo, '```mermaid\npie\n"Done" : 100\n```\n', {})
    # No image embedded, but the build still succeeds (renders as a code block).
    assert _media_names(out) == []


def test_section_bar_shades_headings(tmp_repo):
    out = _build(
        tmp_repo,
        "## Heading\n\nBody.\n",
        {"section_bar": True, "section_bar_color": "#123456"},
    )
    assert 'w:fill="123456"' in _part(out, "word/document.xml")


def test_footer_three_slots_and_page_fields(tmp_repo):
    out = _build(
        tmp_repo,
        "Body.\n",
        {
            "footer_left": "Left",
            "footer_center": "Center",
            "footer_right": "Page {page} of {pages}",
        },
    )
    footer = _part(out, "word/footer1.xml")
    assert "Left" in footer and "Center" in footer
    assert "PAGE" in footer and "NUMPAGES" in footer


def test_nested_lists_are_indented(tmp_repo):
    out = _build(
        tmp_repo,
        "- a\n    - b\n        - c\n",
        {},
    )
    doc_xml = _part(out, "word/document.xml")
    # Deeper levels carry explicit left indents.
    assert 'w:left="720"' in doc_xml and 'w:left="1080"' in doc_xml


# ── PDF↔DOCX page-break / structural parity ─────────────────────────────────


def test_appendix_h2_forces_page_break(tmp_repo):
    # APPENDIX section H2s break the page in both PDF and docx (shared markers).
    # The leading "# Doc" title H1 is what _strip_leading_h1 removes, leaving the
    # in-body "# APPENDIX" heading for the appendix-break pass to find.
    with_appendix = _part(
        _build(tmp_repo, "# Doc\n\n## Intro\n\ntext\n\n# APPENDIX\n\n## A1\n\none\n", {}),
        "word/document.xml",
    )
    assert 'w:type="page"' in with_appendix
    without = _part(_build(tmp_repo, "# Doc\n\n## Intro\n\ntext\n", {}), "word/document.xml")
    assert 'w:type="page"' not in without


def test_explicit_pagebreak_marker(tmp_repo):
    xml = _part(_build(tmp_repo, "a\n\n<!-- pagebreak -->\n\nb\n", {}), "word/document.xml")
    assert 'w:type="page"' in xml


def test_headings_keep_with_next(tmp_repo):
    d = Document(str(_build(tmp_repo, "## Heading\n\nbody text\n", {})))
    h = next(p for p in d.paragraphs if p.text == "Heading")
    assert h.paragraph_format.keep_with_next is True


def test_definition_list_renders(tmp_repo):
    d = Document(str(_build(tmp_repo, "Term A\n:   Definition of A\n", {})))
    term = next(p for p in d.paragraphs if p.text == "Term A")
    assert term.runs[0].bold is True
    dd = next(p for p in d.paragraphs if p.text == "Definition of A")
    assert dd.paragraph_format.left_indent is not None


def test_page_geometry_parser():
    from md_doc.builders.docx import _page_geometry

    a4 = _page_geometry("@page { size: A4; margin: 25mm 20mm 22mm 25mm; }")
    assert (a4["w"], a4["h"]) == (210.0, 297.0)
    assert (a4["top"], a4["right"], a4["bottom"], a4["left"]) == (25.0, 20.0, 22.0, 25.0)
    letter = _page_geometry("@page { size: Letter; margin: 1in; }")
    assert (round(letter["w"], 1), round(letter["h"], 1)) == (215.9, 279.4)
    assert round(letter["top"], 1) == 25.4
    assert _page_geometry(None)["w"] == 210.0  # default A4


def test_cover_matches_pdf_design(tmp_repo):
    # The cover title must be an explicit large bold coloured run (mirroring the
    # PDF's .cover-title) — NOT the built-in serif "Title" style — and the
    # metadata must be colon-free ("Prepared by {author}", not "Prepared by:").
    (tmp_repo / "_pdf-theme.css").write_text(
        "@page { size: A4; margin: 25mm 20mm 22mm 25mm; }\n"
        "body { font-family: Arial; }\n"
        "h1 { color: #1b4f72; }\n"
        "h2 { color: #2e86c1; }\n",
        encoding="utf-8",
    )
    md = "---\ntitle: My Report\n---\n\n# My Report\n\nBody.\n"
    doc_path = tmp_repo / "doc.md"
    doc_path.write_text(md, encoding="utf-8")
    out = tmp_repo / "cover.docx"
    from md_doc.builders.docx import build

    build(
        md,
        {"title": "My Report", "author": "Ada Lovelace", "date": "March 2026"},
        out,
        output_format="docx",
        doc_path=doc_path,
        repo_root=tmp_repo,
    )
    d = Document(str(out))
    from docx.shared import Pt, RGBColor

    title = next(p for p in d.paragraphs if p.text == "My Report" and p.style.name != "Title")
    trun = title.runs[0]
    assert trun.bold is True
    assert trun.font.size == Pt(24)
    assert trun.font.color.rgb == RGBColor(0x1B, 0x4F, 0x72)  # $primary (color_h1)

    # Metadata carries the author/date with NO colon after the label.
    meta = [p.text for p in d.paragraphs]
    assert "Prepared by Ada Lovelace" in meta
    assert "Date March 2026" in meta
    assert not any(t.startswith("Prepared by:") for t in meta)

    # Footer is a text frame anchored ~14mm from the physical page bottom
    # (matching the PDF's .cover-footer { bottom: 14mm }).
    xml = _part(out, "word/document.xml")
    assert "w:framePr" in xml and 'w:vAnchor="page"' in xml and 'w:y="' in xml
    # The cover bar floats at the physical page top (full-bleed, like the
    # PDF's @page cover { margin: 0 }).
    assert "w:tblpPr" in xml and 'w:tblpYSpec="top"' in xml
    # Headers/footers are suppressed on the cover via "different first page".
    assert "<w:titlePg/>" in xml


def test_docx_page_size_matches_theme(tmp_repo):
    # A Letter-sized PDF theme must produce a Letter-sized docx (not hardcoded A4)
    # so the two formats share text width and pagination.
    (tmp_repo / "_pdf-theme.css").write_text(
        "@page { size: Letter; margin: 25mm 20mm 22mm 25mm; }\nbody { font-family: Arial; }\n",
        encoding="utf-8",
    )
    d = Document(str(_build(tmp_repo, "## S\n\nbody\n", {})))
    section = d.sections[0]
    assert round(section.page_width / 36000, 1) == 215.9  # EMU → mm
    assert round(section.page_height / 36000, 1) == 279.4


def test_h1_page_break_before_from_theme(tmp_repo):
    # The PDF theme forces every report-body H1 onto a new page; the docx
    # builder mirrors that — but never on the first content element (a forced
    # break at the top of a page collapses in CSS).
    (tmp_repo / "_pdf-theme.css").write_text(
        "@page { size: A4; margin: 25mm 20mm 22mm 25mm; }\n"
        "body { font-family: Arial; }\n"
        ".report-body h1 { page-break-before: always; break-before: page; }\n",
        encoding="utf-8",
    )
    out = _build(tmp_repo, "# First\n\ntext\n\n# Second\n\nmore\n", {})
    d = Document(str(out))
    first = next(p for p in d.paragraphs if p.text == "First")
    second = next(p for p in d.paragraphs if p.text == "Second")
    assert first.paragraph_format.page_break_before is not True
    assert second.paragraph_format.page_break_before is True


def test_theme_default_footers_render_in_word(tmp_repo):
    # With no footer_* config, the PDF still renders the theme's @bottom-*
    # margin boxes (org name / running date / page numbers). The docx footer
    # must show the same defaults.
    (tmp_repo / "_pdf-theme.css").write_text(
        "@page { size: A4; margin: 25mm 20mm 22mm 25mm;\n"
        '  @bottom-left { content: "Org Name"; font-size: 7.5pt; color: #7f8c9a; }\n'
        "  @bottom-center { content: string(research-date); }\n"
        '  @bottom-right { content: "Page " counter(page) " of " counter(pages); }\n'
        "}\n"
        "body { font-family: Arial; }\n",
        encoding="utf-8",
    )
    md = "---\ntitle: T\n---\n\n## S\n\nbody\n"
    doc_path = tmp_repo / "doc.md"
    doc_path.write_text(md, encoding="utf-8")
    out = tmp_repo / "out.docx"
    from md_doc.builders.docx import build

    build(
        md,
        {"title": "T", "cover_page": False, "date": "1 May 2026"},
        out,
        doc_path=doc_path,
        repo_root=tmp_repo,
    )
    footer = _part(out, "word/footer1.xml")
    assert "Org Name" in footer
    assert "1 May 2026" in footer  # string(research-date) → document date
    assert "PAGE" in footer and "NUMPAGES" in footer
    # An explicitly configured slot still wins over the CSS default.
    build(
        md,
        {"title": "T", "cover_page": False, "date": "1 May 2026", "footer_left": "Custom"},
        out,
        doc_path=doc_path,
        repo_root=tmp_repo,
    )
    footer = _part(out, "word/footer1.xml")
    assert "Custom" in footer and "Org Name" not in footer


def test_no_cover_keeps_h1_in_body(tmp_repo):
    # cover_page: false leaves the leading H1 in the body as a Heading 1 —
    # exactly what the PDF does — rather than a serif "Title" paragraph.
    d = Document(str(_build(tmp_repo, "# My Doc\n\nbody\n", {})))
    h1 = next(p for p in d.paragraphs if p.text == "My Doc")
    assert h1.style.name == "Heading 1"


def test_keep_with_next_excludes_page_break_divs():
    # A forced page-break div must never be swallowed into the PDF's
    # keep-together wrapper — that strands the heading on its own page and can
    # emit an entirely blank page.
    from md_doc.builders.pdf import _keep_heading_with_next

    html = "<h2>Rollout</h2>\n<p>intro</p>\n" '<div class="md-doc-page-break"></div>\n<p>after</p>'
    result = _keep_heading_with_next(html)
    keep = result.split('<div class="keep-with-next">')[1].split("</div>")[0]
    assert "md-doc-page-break" not in keep
