"""Integration tests for the extract CLI command."""

from pathlib import Path

import pytest
from click.testing import CliRunner
from docx import Document

from md_doc.cli import main


def test_extract_docx_to_templates(tmp_path):
    """Extract DOCX to templates/ folder with default naming."""
    # Set up repo structure
    (tmp_path / ".git").mkdir()

    # Create source DOCX
    doc = Document()
    doc.add_heading("Integration Guide", 0)
    doc.add_paragraph("Step 1: Install the package")
    doc.add_paragraph("Step 2: Configure settings")
    source_docx = tmp_path / "integration.docx"
    doc.save(str(source_docx))

    # Create templates folder
    templates_folder = tmp_path / "templates"
    templates_folder.mkdir()

    # Run extract command
    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(source_docx),
            "--dest",
            str(templates_folder),
        ],
    )

    assert result.exit_code == 0, result.output

    # Check output file was created
    output_file = templates_folder / "integration.md"
    assert output_file.exists(), f"Expected {output_file} to exist"

    # Check content was extracted
    content = output_file.read_text()
    assert "Integration Guide" in content
    assert "Step 1" in content


def test_extract_pdf_to_custom_folder(tmp_path):
    """Extract PDF to custom folder with --dest flag."""
    (tmp_path / ".git").mkdir()

    # Create a minimal PDF using reportlab
    try:
        from reportlab.pdfgen import canvas

        pdf_path = tmp_path / "sample.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(100, 750, "Sample PDF Content")
        c.save()
    except ImportError:
        pytest.skip("reportlab not installed for PDF generation")

    # Create custom destination folder
    custom_folder = tmp_path / "snippets"
    custom_folder.mkdir()

    # Run extract command
    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(pdf_path),
            "--dest",
            str(custom_folder),
        ],
    )

    assert result.exit_code == 0, result.output

    # Check output file was created with .md extension
    output_file = custom_folder / "sample.md"
    assert output_file.exists(), f"Expected {output_file} to exist"


def test_extract_unsupported_file_type(tmp_path):
    """Extract should reject unsupported file types."""
    (tmp_path / ".git").mkdir()

    text_file = tmp_path / "readme.txt"
    text_file.write_text("Not a valid input type")

    templates_folder = tmp_path / "templates"
    templates_folder.mkdir()

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(text_file),
            "--dest",
            str(templates_folder),
        ],
    )

    assert result.exit_code != 0
    assert "unsupported file type" in result.output.lower()


def test_extract_file_not_found(tmp_path):
    """Extract should error if source file doesn't exist."""
    (tmp_path / ".git").mkdir()

    templates_folder = tmp_path / "templates"
    templates_folder.mkdir()

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(tmp_path / "nonexistent.pdf"),
            "--dest",
            str(templates_folder),
        ],
    )

    assert result.exit_code != 0
    assert "not found" in result.output.lower()


def test_extract_refuses_overwrite_without_force(tmp_path):
    """Extract should refuse to overwrite an existing file."""
    (tmp_path / ".git").mkdir()

    doc = Document()
    doc.add_heading("Overwrite Test", 0)
    source_docx = tmp_path / "existing.docx"
    doc.save(str(source_docx))

    dest_folder = tmp_path / "out"
    dest_folder.mkdir()
    existing_file = dest_folder / "existing.md"
    existing_file.write_text("original content")

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(source_docx),
            "--dest",
            str(dest_folder),
        ],
    )

    assert result.exit_code != 0
    assert "already exists" in result.output.lower()
    assert existing_file.read_text() == "original content"


def test_extract_overwrites_with_force(tmp_path):
    """Extract should overwrite when --force is used."""
    (tmp_path / ".git").mkdir()

    doc = Document()
    doc.add_heading("Force Overwrite", 0)
    source_docx = tmp_path / "forceme.docx"
    doc.save(str(source_docx))

    dest_folder = tmp_path / "out"
    dest_folder.mkdir()
    existing_file = dest_folder / "forceme.md"
    existing_file.write_text("old content")

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "extract",
            str(source_docx),
            "--dest",
            str(dest_folder),
            "--force",
        ],
    )

    assert result.exit_code == 0, result.output
    assert existing_file.read_text() != "old content"


def test_extract_default_destination_is_templates(tmp_path):
    """If --dest is omitted, default to templates/."""
    (tmp_path / ".git").mkdir()

    # Create source DOCX
    doc = Document()
    doc.add_heading("Auto-save Test", 0)
    source_docx = tmp_path / "auto_test.docx"
    doc.save(str(source_docx))

    # Create templates folder
    templates_folder = tmp_path / "templates"
    templates_folder.mkdir()

    # Run extract without --dest (should use relative templates/ path)
    # Note: CliRunner doesn't change cwd, so we need to verify the behavior works from cwd
    runner = CliRunner()
    with runner.isolated_filesystem():
        # Create templates folder and move file to this context
        import shutil

        Path("templates").mkdir()
        shutil.copy(str(source_docx), "auto_test.docx")

        result = runner.invoke(
            main,
            [
                "extract",
                "auto_test.docx",
            ],
        )

        assert result.exit_code == 0, result.output

        # Check that default templates/ was used
        output_file = Path("templates") / "auto_test.md"
        assert output_file.exists(), f"Expected {output_file} to exist with default dest"
