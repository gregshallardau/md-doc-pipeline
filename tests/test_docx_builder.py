"""Tests for the DOCX/DOTX builder — table column width config."""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from md_doc.builders.docx import build


@pytest.fixture()
def tmp_repo(tmp_path):
    (tmp_path / ".git").mkdir()
    return tmp_path


def _build_docx(tmp_repo: Path, body: str, config: dict, fmt: str = "docx") -> Document:
    md = f"---\ntitle: Test\noutputs: [{fmt}]\n---\n\n{body}"
    out = tmp_repo / f"out.{fmt}"
    build(md, {"title": "Test", "cover_page": False, **config}, out, output_format=fmt)
    return Document(str(out))


def _col_widths(doc: Document) -> list[int]:
    """Return gridCol w values from the first table."""
    tbl = doc.tables[0]._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return []
    return [int(gc.get(qn("w:w"), 0)) for gc in grid.findall(qn("w:gridCol"))]


class TestTableColWidths:
    def test_equal_widths_by_default(self, tmp_repo):
        doc = _build_docx(tmp_repo, "| A | B |\n|---|---|\n| 1 | 2 |\n", {})
        widths = _col_widths(doc)
        assert len(widths) == 2
        assert abs(widths[0] - widths[1]) <= 1  # rounding tolerance of 1 twip

    def test_custom_col_widths_applied(self, tmp_repo):
        doc = _build_docx(tmp_repo, "| A | B |\n|---|---|\n| 1 | 2 |\n",
                          {"table_col_widths": [30, 70]})
        widths = _col_widths(doc)
        assert len(widths) == 2
        # 30:70 ratio — second column should be more than twice the first
        assert widths[1] > widths[0] * 2

    def test_col_widths_sum_to_text_width(self, tmp_repo):
        doc = _build_docx(tmp_repo, "| A | B |\n|---|---|\n| 1 | 2 |\n",
                          {"table_col_widths": [40, 60]})
        widths = _col_widths(doc)
        tbl = doc.tables[0]._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        tblW = tblPr.find(qn("w:tblW"))
        total_declared = int(tblW.get(qn("w:w"), 0))
        assert sum(widths) == total_declared

    def test_mismatched_col_widths_falls_back_to_equal(self, tmp_repo):
        # [30, 70] has 2 values but table has 3 columns — falls back to equal
        doc = _build_docx(tmp_repo, "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |\n",
                          {"table_col_widths": [30, 70]})
        widths = _col_widths(doc)
        assert len(widths) == 3
        assert abs(widths[0] - widths[1]) <= 1
        assert abs(widths[1] - widths[2]) <= 1

    def test_cell_widths_match_grid(self, tmp_repo):
        doc = _build_docx(tmp_repo, "| A | B |\n|---|---|\n| 1 | 2 |\n",
                          {"table_col_widths": [30, 70]})
        grid_widths = _col_widths(doc)
        table = doc.tables[0]
        for row in table.rows:
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.find(qn("w:tcPr"))
                tcW = tcPr.find(qn("w:tcW"))
                assert int(tcW.get(qn("w:w"))) == grid_widths[c_idx]

    def test_only_one_tblgrid(self, tmp_repo):
        """Fixed layout must not have duplicate tblGrid elements."""
        doc = _build_docx(tmp_repo, "| A | B |\n|---|---|\n| 1 | 2 |\n", {})
        tbl = doc.tables[0]._tbl
        grids = tbl.findall(qn("w:tblGrid"))
        assert len(grids) == 1

    def test_short_rows_all_cells_get_tcW(self, tmp_repo):
        """Rows with fewer cells than max_cols must still have explicit tcW on all cells."""
        # Row 1 has 3 cols, row 2 has 2 cols — the 3rd cell in row 2 must still get tcW
        body = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |\n| x | y |\n"
        doc = _build_docx(tmp_repo, body, {})
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                assert tcPr is not None
                tcW = tcPr.find(qn("w:tcW"))
                assert tcW is not None, "Every cell must have explicit tcW for fixed layout"
                assert int(tcW.get(qn("w:w"))) > 0


class TestColWidthsComment:
    def test_comment_sets_widths_for_that_table(self, tmp_repo):
        body = "<!-- col-widths: 30, 70 -->\n| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = _build_docx(tmp_repo, body, {})
        widths = _col_widths(doc)
        assert len(widths) == 2
        assert widths[1] > widths[0] * 2

    def test_comment_only_applies_to_next_table(self, tmp_repo):
        body = (
            "<!-- col-widths: 30, 70 -->\n| A | B |\n|---|---|\n| 1 | 2 |\n\n"
            "| C | D |\n|---|---|\n| 3 | 4 |\n"
        )
        doc = _build_docx(tmp_repo, body, {})

        def grid_widths(table_idx):
            tbl = doc.tables[table_idx]._tbl
            grid = tbl.find(qn("w:tblGrid"))
            return [int(gc.get(qn("w:w"), 0)) for gc in grid.findall(qn("w:gridCol"))]

        w1 = grid_widths(0)
        w2 = grid_widths(1)
        # First table: 30/70 split
        assert w1[1] > w1[0] * 2
        # Second table: equal (no comment)
        assert abs(w2[0] - w2[1]) <= 1

    def test_comment_overrides_config_widths(self, tmp_repo):
        body = "<!-- col-widths: 60, 40 -->\n| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = _build_docx(tmp_repo, body, {"table_col_widths": [30, 70]})
        widths = _col_widths(doc)
        # Comment says 60/40, config says 30/70 — comment wins, so first >= second
        assert widths[0] > widths[1]


class TestMergeFieldsInTableCells:
    def _dotx_xml(self, path: Path) -> str:
        import zipfile
        with zipfile.ZipFile(path) as zf:
            return zf.read("word/document.xml").decode("utf-8")

    def test_form_field_in_cell_creates_word_field(self, tmp_repo):
        body = "| Label | Value |\n|---|---|\n| Client | [[contact_name]] |\n"
        out = tmp_repo / "out.dotx"
        from md_doc.builders.docx import build
        build("---\ntitle: T\n---\n\n" + body,
              {"title": "T", "cover_page": False}, out, output_format="dotx")
        xml = self._dotx_xml(out)
        assert "fldChar" in xml or "bookmarkStart" in xml

    def test_form_field_not_left_as_literal_text(self, tmp_repo):
        body = "| Label | Value |\n|---|---|\n| Client | [[contact_name]] |\n"
        out = tmp_repo / "out.dotx"
        from md_doc.builders.docx import build
        build("---\ntitle: T\n---\n\n" + body,
              {"title": "T", "cover_page": False}, out, output_format="dotx")
        xml = self._dotx_xml(out)
        assert "[[contact_name]]" not in xml

    def test_merge_field_in_cell(self, tmp_repo):
        body = "| Label | Value |\n|---|---|\n| Client | [[contact_name]] |\n"
        out = tmp_repo / "out.dotx"
        from md_doc.builders.docx import build
        build("---\ntitle: T\n---\n\n" + body,
              {"title": "T", "cover_page": False, "dotx_field_type": "merge"},
              out, output_format="dotx")
        xml = self._dotx_xml(out)
        assert "MERGEFIELD" in xml
