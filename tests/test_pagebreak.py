"""Tests for the ``<!-- pagebreak -->`` marker.

Covers both the PDF preprocessor (markdown → div) and the DOCX walker
(div → real Word page break).
"""

import pytest

from md_doc.builders.pdf import _inject_page_breaks

# ── PDF: markdown preprocessor ────────────────────────────────────────────────


class TestInjectPageBreaks:
    def test_basic_substitution(self):
        md = "Top\n\n<!-- pagebreak -->\n\nBottom"
        result = _inject_page_breaks(md)
        assert '<div class="md-doc-page-break"></div>' in result
        assert "<!-- pagebreak -->" not in result

    def test_case_insensitive(self):
        for marker in ("<!-- PAGEBREAK -->", "<!-- PageBreak -->", "<!--pagebreak-->"):
            assert '<div class="md-doc-page-break"></div>' in _inject_page_breaks(marker)

    def test_multiple_markers(self):
        md = "A\n<!-- pagebreak -->\nB\n<!-- pagebreak -->\nC"
        result = _inject_page_breaks(md)
        assert result.count('<div class="md-doc-page-break"></div>') == 2

    def test_no_marker_unchanged(self):
        md = "Just plain content with no marker"
        assert _inject_page_breaks(md) == md

    def test_other_html_comments_untouched(self):
        md = "<!-- not a pagebreak -->\nbody"
        result = _inject_page_breaks(md)
        assert "<!-- not a pagebreak -->" in result
        assert '<div class="md-doc-page-break"></div>' not in result

    def test_div_surrounded_by_blank_lines(self):
        """Markdown needs blank lines around block-level HTML."""
        md = "para 1\n<!-- pagebreak -->\npara 2"
        result = _inject_page_breaks(md)
        # Should have blank lines around the div so markdown parses it as block
        assert '\n\n<div class="md-doc-page-break"></div>\n\n' in result


# ── PDF: end-to-end build smoke test ──────────────────────────────────────────


class TestPdfBuildPagebreak:
    """Verify the marker survives the full PDF build pipeline."""

    def test_builds_without_error(self, tmp_path):
        pytest.importorskip("weasyprint")
        from md_doc.builders.pdf import build as build_pdf

        (tmp_path / ".git").mkdir()
        out = tmp_path / "out.pdf"

        md = "---\ntitle: Test\n---\n\n" "First page\n\n" "<!-- pagebreak -->\n\n" "Second page\n"

        build_pdf(md, {"title": "Test", "cover_page": False}, out, repo_root=tmp_path)
        assert out.exists()
        assert out.stat().st_size > 0


# ── DOCX: walker emits a real page break ──────────────────────────────────────


class TestDocxPagebreak:
    def test_marker_emits_word_page_break(self, tmp_path):
        from md_doc.builders.docx import build as build_docx

        (tmp_path / ".git").mkdir()
        out = tmp_path / "out.docx"

        md = "---\ntitle: Test\n---\n\n" "Before break\n\n" "<!-- pagebreak -->\n\n" "After break\n"

        build_docx(md, {"title": "Test"}, out, repo_root=tmp_path)
        assert out.exists()

        # Re-open the document and verify a page-break run was inserted.
        # python-docx represents page breaks as <w:br w:type="page"/>.
        from docx import Document

        doc = Document(str(out))
        xml = "\n".join(p._p.xml for p in doc.paragraphs)
        assert 'w:type="page"' in xml or 'type="page"' in xml

    def test_no_marker_no_page_break(self, tmp_path):
        from md_doc.builders.docx import build as build_docx
        from docx import Document

        (tmp_path / ".git").mkdir()
        out = tmp_path / "no_break.docx"

        md = "---\ntitle: Test\n---\n\nJust one page of content.\n"

        build_docx(md, {"title": "Test"}, out, repo_root=tmp_path)

        doc = Document(str(out))
        xml = "\n".join(p._p.xml for p in doc.paragraphs)
        # Cover page is not added for plain .docx, so no page break should appear.
        assert 'w:type="page"' not in xml
