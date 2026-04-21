"""
Word merge-template (.dotx) builder.

Converts rendered Markdown to a Word Template file (.dotx) suitable for
either direct fill-in or mail merge, depending on ``dotx_field_type``.

Field syntax
------------
Use ``[[field_name]]`` in Markdown source. This is intentionally distinct
from Jinja2 ``{{ }}`` so both can coexist in the same document:

    Dear [[contact_name]],

    This is version {{ version }} of our proposal for [[client]].

- ``{{ version }}``     — resolved at build time from _meta.yml / frontmatter
- ``[[contact_name]]``  — becomes a Word field in the .dotx (type controlled
                          by ``dotx_field_type``)

Config keys consumed
--------------------
  title            — cover page heading (may contain [[fields]])
  author           — author name (may contain [[fields]])
  date             — date string (may contain [[fields]])
  product          — product / client subtitle on cover page
  cover_page       — bool (default: true); set false to omit the cover page
  dotx_field_type  — "form" (default) or "merge"

    "form"   — Word Text Form Field with Bookmark = field_name and Fill-in
               enabled. The template is directly fillable in Word without a
               mail merge data source (tab through fields, type, save).

    "merge"  — Classic Word MERGEFIELD (``«field_name»``). Requires a data
               source to be attached and a mail merge to be run before the
               fields fill.

``{% include %}`` fragments work normally; the renderer processes them
before this builder runs.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path
from typing import Any

import markdown
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .docx import (
    _DocxBuilder,
    _MD_EXTENSIONS,
    _extract_title,
    _strip_frontmatter,
    _strip_leading_h1,
)

_MERGE_RE = re.compile(r"\[\[(\w+)\]\]")


# ---------------------------------------------------------------------------
# Field helpers — MERGEFIELD and Text Form Field
# ---------------------------------------------------------------------------


def _insert_merge_field(
    paragraph: Any,
    field_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Append a Word MERGEFIELD for *field_name* to *paragraph*."""
    run = paragraph.add_run()
    run.bold = bold
    run.italic = italic
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" MERGEFIELD {field_name} "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run(f"«{field_name}»")
    run.bold = bold
    run.italic = italic

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _insert_form_field(
    paragraph: Any,
    field_name: str,
    bookmark_id: int,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Append a Word Text Form Field with Bookmark=*field_name* to *paragraph*.

    The bookmark name IS the variable name — downstream systems (and Word's own
    Fill-in dialog) identify the field by it. Fill-in is enabled so the template
    is directly fillable in Word without a mail merge data source.
    """
    p = paragraph._p

    # Bookmark start wraps the entire field char sequence
    bk_start = OxmlElement("w:bookmarkStart")
    bk_start.set(qn("w:id"), str(bookmark_id))
    bk_start.set(qn("w:name"), field_name)
    p.append(bk_start)

    # fldChar begin — carries the ffData (form field metadata)
    run = paragraph.add_run()
    run.bold = bold
    run.italic = italic
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    ff_data = OxmlElement("w:ffData")
    ff_name = OxmlElement("w:name")
    ff_name.set(qn("w:val"), field_name)
    ff_data.append(ff_name)
    ff_data.append(OxmlElement("w:enabled"))
    ff_calc = OxmlElement("w:calcOnExit")
    ff_calc.set(qn("w:val"), "0")
    ff_data.append(ff_calc)
    ff_data.append(OxmlElement("w:textInput"))
    fld_begin.append(ff_data)
    run._r.append(fld_begin)

    # FORMTEXT instruction
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " FORMTEXT "
    run._r.append(instr)

    # separate
    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # display text shown in Word before the field is filled
    run = paragraph.add_run(f"«{field_name}»")
    run.bold = bold
    run.italic = italic

    # end
    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # Bookmark end
    bk_end = OxmlElement("w:bookmarkEnd")
    bk_end.set(qn("w:id"), str(bookmark_id))
    p.append(bk_end)


# ---------------------------------------------------------------------------
# HTML → docx walker (field-type aware)
# ---------------------------------------------------------------------------


class _DotxBuilder(_DocxBuilder):
    """
    Extends _DocxBuilder to convert ``[[field]]`` markers to Word fields.

    Field type is controlled by *field_type*:
    - "form"  — Text Form Fields with Bookmark (directly fillable in Word)
    - "merge" — Classic MERGEFIELDs (require a mail merge data source)

    Only _add_text and _flush_table are overridden; all structural parsing
    (headings, lists, blockquotes, etc.) is inherited unchanged.
    """

    def __init__(self, doc: Any, *, field_type: str = "form") -> None:
        super().__init__(doc)
        self._field_type = field_type
        self._bookmark_id = 0

    def _write_text(
        self,
        paragraph: Any,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        code: bool = False,
    ) -> None:
        """Write *text* to *paragraph*, converting ``[[field]]`` to Word fields."""
        if not text:
            return
        parts = _MERGE_RE.split(text)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    run = paragraph.add_run(part)
                    run.bold = bold
                    run.italic = italic
                    if code:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
            else:
                if self._field_type == "merge":
                    _insert_merge_field(paragraph, part, bold=bold, italic=italic)
                else:
                    _insert_form_field(paragraph, part, self._bookmark_id, bold=bold, italic=italic)
                    self._bookmark_id += 1

    def _add_text(self, text: str) -> None:
        if not text:
            return
        self._write_text(
            self._current_para(),
            text,
            bold=self._bold,
            italic=self._italic,
            code=self._in_code,
        )

    def _flush_table(self) -> None:
        rows = getattr(self, "_table_rows", [])
        if not rows:
            return
        max_cols = max(len(r) for r in rows)
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx, (is_header, text) in enumerate(row):
                if c_idx >= max_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                self._write_text(cell.paragraphs[0], text.strip(), bold=is_header)

        self._paragraph = None


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------


def _add_cover_page(doc: Document, config: dict[str, Any], builder: "_DotxBuilder") -> None:
    """Insert a cover page section followed by a page break.

    All text values may contain ``[[field]]`` markers — rendered via *builder*
    so the correct field type (form/merge) and bookmark ID counter are used.
    """
    title = config.get("title", "")
    author = config.get("author", "")
    date = config.get("date", "")
    product = config.get("product", "")

    title_para = doc.add_paragraph(style="Title")
    builder._write_text(title_para, title or "«Document Title»")

    if product:
        sub_para = doc.add_paragraph(style="Subtitle")
        builder._write_text(sub_para, product)

    if author or date:
        doc.add_paragraph()  # spacer
        if author:
            meta = doc.add_paragraph()
            meta.add_run("Prepared by: ").bold = True
            builder._write_text(meta, author)
        if date:
            meta = doc.add_paragraph()
            meta.add_run("Date: ").bold = True
            builder._write_text(meta, date)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# .dotx content-type patch
# ---------------------------------------------------------------------------


def _patch_to_dotx(path: Path) -> None:
    """
    Re-write *path* with the Word Template content type.

    python-docx always saves with the .docx content type. A .dotx differs
    only in one attribute inside ``[Content_Types].xml`` inside the ZIP.
    """
    tmp = path.with_suffix(".tmp")
    shutil.move(str(path), str(tmp))
    try:
        with zipfile.ZipFile(tmp, "r") as zin:
            with zipfile.ZipFile(path, "w") as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        data = data.replace(
                            b"wordprocessingml.document.main+xml",
                            b"wordprocessingml.template.main+xml",
                        )
                    zout.writestr(item, data)
    except Exception:
        # Restore original on failure
        shutil.move(str(tmp), str(path))
        raise
    finally:
        if tmp.exists():
            tmp.unlink()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build(
    rendered_md: str,
    config: dict[str, Any],
    out_path: Path,
    *,
    doc_path: Path | None = None,
) -> None:
    """
    Convert rendered Markdown to a .dotx Word merge template.

    Parameters
    ----------
    rendered_md:
        Jinja2-rendered Markdown string. ``[[field]]`` markers are left
        intact by the renderer and converted to MERGEFIELDs here.
    config:
        Merged config dict from load_config().
    out_path:
        Destination path for the generated .dotx file.
    doc_path:
        Source .md path (unused currently, reserved for future use).
    """
    out_path = Path(out_path).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    field_type = str(config.get("dotx_field_type", "form")).lower()
    if field_type not in ("form", "merge"):
        field_type = "form"

    body = _strip_frontmatter(rendered_md)
    title: str = config.get("title") or _extract_title(body) or out_path.stem
    author: str = config.get("author", "")
    cover_page: bool = bool(config.get("cover_page", True))

    if cover_page:
        body = _strip_leading_h1(body)

    # Markdown → HTML
    md_engine = markdown.Markdown(extensions=_MD_EXTENSIONS)
    html = md_engine.convert(body)

    # Build document
    doc = Document()

    props = doc.core_properties
    props.title = re.sub(r"\[\[\w+\]\]", "", title).strip()
    if author:
        props.author = re.sub(r"\[\[\w+\]\]", "", author).strip()

    builder = _DotxBuilder(doc, field_type=field_type)

    if cover_page:
        _add_cover_page(doc, {**config, "title": title}, builder)

    builder.feed(html)

    # Save as .docx first (python-docx limitation), then patch content type
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _patch_to_dotx(out_path)
