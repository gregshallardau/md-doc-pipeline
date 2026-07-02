"""
python-docx builder for .docx and .dotx output.

Converts rendered Markdown to a Word document.  When ``output_format="dotx"``
the builder also converts ``[[field_name]]`` markers to Word fields and patches
the saved ZIP's content type so Word opens it as a template.

Public API
----------
    build(rendered_md, config, out_path, *, doc_path, repo_root, output_format)

Field syntax (.dotx only)
-------------------------
Use ``[[field_name]]`` in Markdown source alongside Jinja2 ``{{ }}``:

    Dear [[contact_name]],          # Word field in .dotx
    This is version {{ version }}.  # resolved at build time

``dotx_field_type`` config key controls field type:
  "form"  (default) — Text Form Fields with Bookmark; fillable in Word without
                       a mail merge data source.
  "merge"           — Classic MERGEFIELD (``«field_name»``); requires a data
                       source and a mail merge run.
"""

from __future__ import annotations

import datetime
import logging
import re
import shutil
import zipfile
from io import BytesIO
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

from ..docx_theme import (
    _apply_font_name,
    _hex_to_rgb,
    apply_theme_to_doc,
    patch_docx_theme_fonts,
    resolve_docx_theme,
    set_cell_shading,
    set_para_shading,
)
from .pdf import _inject_appendix_breaks, _inject_page_breaks
from ._assets import (
    _EMU_PER_PX,
    _MERMAID_IMG_RE,
    _render_mermaid_to_images,
    _resolve_asset,
    _svg_to_png,
)

logger = logging.getLogger(__name__)

# Re-exported for backwards compatibility (tests/other modules import these
# from docx); the implementations now live in ._assets.
__all__ = [
    "build",
    "_EMU_PER_PX",
    "_MERMAID_IMG_RE",
    "_render_mermaid_to_images",
    "_resolve_asset",
    "_svg_to_png",
]


# Markdown extensions (consistent with pdf builder)
_MD_EXTENSIONS = [
    "tables",
    "fenced_code",
    "footnotes",
    "def_list",
    "abbr",
    "attr_list",
    "md_in_html",
    "toc",
]

_MERGE_RE = re.compile(r"\[\[(\w+)\]\]")


# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------


def _insert_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Append a clickable hyperlink to *paragraph*.

    Display text is ``text (url)`` when text differs from url, or just ``url``
    when they're the same (e.g. bare URL links). This keeps the URL visible
    in printed output.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    display = text.strip()
    url_clean = url.strip()
    if display and display != url_clean:
        display = f"{display} ({url_clean})"
    else:
        display = url_clean

    part = paragraph.part
    r_id = part.relate_to(url_clean, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = display
    if display.startswith(" ") or display.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Field helpers — MERGEFIELD and Text Form Field
# ---------------------------------------------------------------------------


def _insert_merge_field(
    paragraph: Any,
    field_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Append a Word MERGEFIELD for *field_name* to *paragraph*."""
    run = paragraph.add_run()
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" MERGEFIELD {field_name} "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run(f"«{field_name}»")
    if bold:
        run.bold = True
    if italic:
        run.italic = True

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _insert_form_field(
    paragraph: Any,
    field_name: str,
    bookmark_id: int,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Append a Word Text Form Field named *field_name* to *paragraph*.

    The field name is stored in ffData/name and is directly fillable in Word
    without a mail merge data source.  Bookmarks are intentionally omitted —
    they are not needed for fill-in use and cause Word to render extra visual
    line breaks when multiple fields appear consecutively in the same paragraph.
    """
    run = paragraph.add_run()
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    ff_data = OxmlElement("w:ffData")
    ff_name = OxmlElement("w:name")
    ff_name.set(qn("w:val"), field_name)
    ff_data.append(ff_name)
    ff_data.append(OxmlElement("w:enabled"))
    ff_calc = OxmlElement("w:calcOnExit")
    ff_calc.set(qn("w:val"), "0")
    ff_data.append(ff_calc)
    ff_data.append(OxmlElement("w:textInput"))
    fld_begin.append(ff_data)
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " FORMTEXT "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run(f"«{field_name}»")
    if bold:
        run.bold = True
    if italic:
        run.italic = True

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------


def _set_cell_bottom_border(cell: Any, color: str = "d5d8dc", pt: float = 0.5) -> None:
    """Add a bottom border to a table cell. ``pt`` is the border thickness in points."""
    sz = str(max(1, round(pt * 8)))  # Word sz unit = 1/8 pt
    fill = color.lstrip("#").upper()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    bottom = tcBorders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        tcBorders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), fill)


def _clear_table_borders(table: Any) -> None:
    """Explicitly set all table-level borders to none so cell borders control the look."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _render_cell_html(
    paragraph: Any,
    html: str,
    theme: dict,
    write_text: Any,
    *,
    bold_override: bool = False,
) -> None:
    """Parse the inner HTML of a table cell and write runs into *paragraph*.

    Handles inline tags: strong/b (bold), em/i (italic), code (monospace).
    Text (including ``[[field]]`` markers) is written via *write_text*, which
    is the builder's ``_write_text`` method — ensuring field conversion and
    bookmark tracking work identically to body text.
    """
    from html.parser import HTMLParser as _HP

    class _CellParser(_HP):
        def __init__(self) -> None:
            super().__init__()
            self._bold = bold_override
            self._italic = False
            self._code = False

        def handle_starttag(self, tag: str, attrs: list) -> None:
            tag = tag.lower()
            if tag in ("strong", "b"):
                self._bold = True
            elif tag in ("em", "i"):
                self._italic = True
            elif tag == "code":
                self._code = True
            elif tag == "br":
                br_run = paragraph.add_run()
                br_run._r.append(OxmlElement("w:br"))

        def handle_endtag(self, tag: str) -> None:
            tag = tag.lower()
            if tag in ("strong", "b"):
                self._bold = bold_override  # back to cell default
            elif tag in ("em", "i"):
                self._italic = False
            elif tag == "code":
                self._code = False

        def handle_data(self, data: str) -> None:
            if not data:
                return
            write_text(paragraph, data, bold=self._bold, italic=self._italic, code=self._code)

    _CellParser().feed(html)


# ---------------------------------------------------------------------------
# HTML → docx walker
# ---------------------------------------------------------------------------


class _DocxBuilder(HTMLParser):
    """
    Walk an HTML fragment and populate a python-docx Document.

    Handles: h1–h4, p, ul/ol/li, table/thead/tbody/tr/th/td,
             pre/code, blockquote, strong/b, em/i, hr, br.

    When *field_type* is ``"form"`` or ``"merge"``, ``[[field_name]]``
    markers in text are converted to the appropriate Word field type
    instead of being written as literal text.
    """

    def __init__(
        self,
        doc: Document,
        theme: dict[str, Any] | None = None,
        field_type: str | None = None,
        body_text_align: str | None = None,
        table_col_widths: list[float] | None = None,
        *,
        mermaid_images: list[tuple[bytes, int, int]] | None = None,
        doc_path: Path | None = None,
        repo_root: Path | None = None,
        section_bar: dict[str, Any] | None = None,
    ) -> None:
        super().__init__()
        self.convert_charrefs = True
        self.doc = doc
        self._theme: dict[str, Any] = theme or {}
        self._field_type = field_type  # None | "form" | "merge"
        self._body_text_align = body_text_align  # default alignment for Normal paragraphs
        self._table_col_widths = table_col_widths  # e.g. [30, 70] — relative column widths
        self._mermaid_images = mermaid_images or []
        self._doc_path = doc_path
        self._repo_root = repo_root
        self._section_bar = section_bar  # None or parsed section_bar config

        apply_theme_to_doc(self.doc, self._theme)

        # State tracking
        self._paragraph = None
        self._run = None
        self._bold = False
        self._italic = False
        self._in_pre = False
        self._in_code = False
        self._in_blockquote = False
        self._list_stack: list[str] = []
        self._list_counters: list[int] = []

        # Table state — cells store raw inner HTML to preserve inline markup
        self._in_table = False
        self._in_cell = False  # True while cursor is inside a <th> or <td>
        self._in_th = False
        self._table_rows: list[list[tuple[bool, str]]] = []
        self._current_row: list[tuple[bool, str]] = []
        self._current_cell_html = ""
        # Set by <!-- col-widths: 30, 70 --> comments; consumed by the next table
        self._next_table_col_widths: list[float] | None = None
        self._active_table_col_widths: list[float] | None = None

        self._pre_text = ""
        self._tag_stack: list[str] = []

        # Field-mode state
        self._bookmark_id = 0

        # Set to True after a <br> element is written so the leading \n of the
        # next handle_data call (which is just HTML formatting whitespace after
        # the <br> tag, not a meaningful line break) is stripped rather than
        # converted to an extra <w:br/>.
        self._last_was_br = False

        # Alignment context stack — pushed/popped by <div style="text-align: ...">
        self._alignment_stack: list[str | None] = []

        # Hyperlink state — set while inside <a href="...">
        self._current_href: str | None = None
        self._link_text_buf: str = ""

        # Section-bar state — the heading tag currently wearing a bar (or None)
        self._section_bar_active_tag: str | None = None

        # Body-length baseline for page-break-before headings.  A heading that
        # is the very first content element must not force a break (matching
        # CSS, where a forced break at the top of a page collapses).  build()
        # re-baselines this after the cover page is added.
        self.mark_content_start()

    def mark_content_start(self) -> None:
        """Record the current body length as the start of report content."""
        self._body_baseline = len(self.doc.element.body)

    # ------------------------------------------------------------------
    # Alignment helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_text_align(attrs: list[tuple[str, str | None]]) -> str | None:
        """Extract text-align value from a style attribute, e.g. 'justify'."""
        for name, value in attrs:
            if name == "style" and value:
                for part in value.split(";"):
                    part = part.strip()
                    if part.lower().startswith("text-align:"):
                        return part.split(":", 1)[1].strip().lower()
        return None

    @staticmethod
    def _to_word_alignment(align: str | None) -> Any:
        _map = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        return _map.get(align or "")

    def _effective_alignment(self, inline_align: str | None = None) -> Any:
        """Return the Word alignment constant for the current context."""
        align = inline_align
        if align is None:
            # Walk stack top-to-bottom for nearest div override
            for a in reversed(self._alignment_stack):
                if a is not None:
                    align = a
                    break
        if align is None:
            align = self._body_text_align
        return self._to_word_alignment(align)

    # ------------------------------------------------------------------
    # Section bar helpers (mirror pdf._build_section_bar_style)
    # ------------------------------------------------------------------

    def _apply_section_bar_start(self, tag: str) -> None:
        """Apply the section-bar background/border to a freshly created heading."""
        self._section_bar_active_tag = None
        sb = self._section_bar
        if not sb or tag not in sb["headings"] or self._paragraph is None:
            return
        self._section_bar_active_tag = tag
        para = self._paragraph
        if sb["text_on_bar"]:
            set_para_shading(para, sb["color"].lstrip("#"))
            # Snug the bar around the text (CSS padding: 6pt 12pt).
            pf = para.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            ind = para._p.get_or_add_pPr().get_or_add_ind()
            ind.set(qn("w:left"), str(int(12 * 20)))
            ind.set(qn("w:right"), str(int(12 * 20)))
        else:
            # border-top variant
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), str(max(1, round(4 * 8))))  # 4pt
            top.set(qn("w:space"), "6")
            top.set(qn("w:color"), sb["color"].lstrip("#").upper())
            pBdr.append(top)
            pPr.append(pBdr)

    def _apply_section_bar_runs(self) -> None:
        """Colour the heading's runs white once its text has been written."""
        sb = self._section_bar
        if not sb or not sb["text_on_bar"] or self._paragraph is None:
            return
        r, g, b = _hex_to_rgb(sb["text_color"])
        for run in self._paragraph.runs:
            run.font.color.rgb = RGBColor(r, g, b)

    # ------------------------------------------------------------------
    # Paragraph helpers
    # ------------------------------------------------------------------

    def _new_para(self, style: str = "Normal") -> None:
        self._paragraph = self.doc.add_paragraph(style=style)
        self._run = None

    def _current_para(self) -> Any:
        if self._paragraph is None:
            self._paragraph = self.doc.add_paragraph()
        return self._paragraph

    def _write_text(
        self,
        paragraph: Any,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        code: bool = False,
    ) -> None:
        """Write *text* to *paragraph*, handling ``[[field]]`` markers when
        field_type is set, and converting bare ``\\n`` to Word line breaks.
        """
        if not text:
            return

        if self._field_type:
            # Field-aware path: split on [[field]] markers
            parts = _MERGE_RE.split(text)
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Literal text segment — split on \n for line breaks
                    if part:
                        lines = part.split("\n")
                        for j, line in enumerate(lines):
                            if j > 0:
                                br_run = paragraph.add_run()
                                br_run._r.append(OxmlElement("w:br"))
                            if line:
                                run = paragraph.add_run(line)
                                if bold:
                                    run.bold = True
                                    col = self._theme.get("color_strong")
                                    if col:
                                        r, g, b = _hex_to_rgb(col)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                if italic:
                                    run.italic = True
                                    if not bold:
                                        col = self._theme.get("color_em")
                                        if col:
                                            r, g, b = _hex_to_rgb(col)
                                            run.font.color.rgb = RGBColor(r, g, b)
                                if code:
                                    run.font.name = self._theme.get("font_code", "Courier New")
                                    run.font.size = Pt(self._theme.get("font_size_code", 9.0))
                                    col = self._theme.get("color_code")
                                    if col:
                                        r, g, b = _hex_to_rgb(col)
                                        run.font.color.rgb = RGBColor(r, g, b)
                else:
                    # Field marker
                    if self._field_type == "merge":
                        _insert_merge_field(paragraph, part, bold=bold, italic=italic)
                    else:
                        _insert_form_field(
                            paragraph, part, self._bookmark_id, bold=bold, italic=italic
                        )
                        self._bookmark_id += 1
        else:
            # Plain text path: split on \n for line breaks
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    br_run = paragraph.add_run()
                    br_run._r.append(OxmlElement("w:br"))
                if line:
                    run = paragraph.add_run(line)
                    if bold:
                        run.bold = True
                        col = self._theme.get("color_strong")
                        if col:
                            r, g, b = _hex_to_rgb(col)
                            run.font.color.rgb = RGBColor(r, g, b)
                    if italic:
                        run.italic = True
                        if not bold:
                            col = self._theme.get("color_em")
                            if col:
                                r, g, b = _hex_to_rgb(col)
                                run.font.color.rgb = RGBColor(r, g, b)
                    if code:
                        run.font.name = self._theme.get("font_code", "Courier New")
                        run.font.size = Pt(self._theme.get("font_size_code", 9.0))
                        col = self._theme.get("color_code")
                        if col:
                            r, g, b = _hex_to_rgb(col)
                            run.font.color.rgb = RGBColor(r, g, b)

    def _add_text(self, text: str) -> None:
        if not text:
            return
        if self._last_was_br:
            text = text.lstrip("\n")
            self._last_was_br = False
        if self._current_href is not None:
            self._link_text_buf += text
            return
        if self._paragraph is None and not text.strip():
            return
        self._write_text(
            self._current_para(),
            text,
            bold=self._bold,
            italic=self._italic,
            code=self._in_code,
        )

    # ------------------------------------------------------------------
    # HTMLParser callbacks
    # ------------------------------------------------------------------

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self._tag_stack.append(tag)
        tag = tag.lower()

        # Inside a cell, collect all inline tags as raw HTML instead of processing them
        if self._in_cell:
            attr_str = ""
            for k, v in attrs:
                attr_str += f' {k}="{v}"' if v is not None else f" {k}"
            self._current_cell_html += f"<{tag}{attr_str}>"
            return

        if tag in ("h1", "h2", "h3", "h4"):
            self._new_para(f"Heading {int(tag[1])}")
            inline_align = self._parse_text_align(attrs)
            word_align = self._effective_alignment(inline_align)
            if self._paragraph is not None:
                if word_align is not None:
                    self._paragraph.alignment = word_align
                # Keep the heading on the same page as the content that follows
                # (mirrors the PDF builder's keep-heading-with-next behaviour so
                # headings don't strand at a page bottom in one format only).
                self._paragraph.paragraph_format.keep_with_next = True
                # Forced page break before this heading level (mirrors the PDF
                # theme's `.report-body h1 { page-break-before: always }`).
                # Skipped for the first content element — a forced break at the
                # top of a page collapses in CSS, so Word must not add one.
                if self._theme.get(f"page_break_before_{tag}") and (
                    len(self.doc.element.body) - self._body_baseline > 1
                ):
                    self._paragraph.paragraph_format.page_break_before = True
            self._apply_section_bar_start(tag)

        elif tag == "p":
            self._new_para("Normal")
            inline_align = self._parse_text_align(attrs)
            word_align = self._effective_alignment(inline_align)
            if word_align is not None and self._paragraph is not None:
                self._paragraph.alignment = word_align
            if self._in_blockquote:
                # Apply left border accent from theme
                bq_color = self._theme.get("blockquote_border_color")
                bq_pt = self._theme.get("blockquote_border_pt", 3.0)
                if bq_color:
                    pPr = self._paragraph._p.get_or_add_pPr()
                    pBdr = OxmlElement("w:pBdr")
                    left = OxmlElement("w:left")
                    left.set(qn("w:val"), "single")
                    left.set(qn("w:sz"), str(max(1, round(bq_pt * 8))))
                    left.set(qn("w:space"), "12")
                    left.set(qn("w:color"), bq_color.lstrip("#").upper())
                    pBdr.append(left)
                    pPr.append(pBdr)
                # Indentation to match CSS padding-left
                ind = self._paragraph._p.get_or_add_pPr().get_or_add_ind()
                ind.set(qn("w:left"), str(int(10 * 20)))  # 10pt indent

        elif tag == "div":
            self._alignment_stack.append(self._parse_text_align(attrs))
            classes = (dict(attrs).get("class") or "").split()
            # Both the explicit <!-- pagebreak --> marker and the APPENDIX
            # auto-break marker (shared with the PDF builder) emit a real Word
            # page break so the two formats break at the same points.
            if "md-doc-page-break" in classes or "appendix-template-break" in classes:
                self.doc.add_page_break()

        elif tag == "a":
            self._current_href = dict(attrs).get("href") or ""
            self._link_text_buf = ""

        elif tag in ("ul", "ol"):
            self._list_stack.append(tag)
            self._list_counters.append(0)

        elif tag == "li":
            if self._list_stack:
                kind = self._list_stack[-1]
                if kind == "ul":
                    self._paragraph = self.doc.add_paragraph(style="List Bullet")
                else:
                    self._list_counters[-1] += 1
                    self._paragraph = self.doc.add_paragraph(style="List Number")
            else:
                self._new_para("List Bullet")
            # Indent nested list items so depth is visible (the base list styles
            # only indent one level, matching CSS nested-list indentation).
            depth = max(len(self._list_stack), 1)
            if depth > 1 and self._paragraph is not None:
                self._paragraph.paragraph_format.left_indent = Pt(18 * depth)

        elif tag == "dt":
            # Definition-list term — a bold Normal paragraph.
            self._new_para("Normal")
            self._bold = True

        elif tag == "dd":
            # Definition-list description — indented Normal paragraph.
            self._new_para("Normal")
            if self._paragraph is not None:
                self._paragraph.paragraph_format.left_indent = Pt(18)

        elif tag == "pre":
            self._in_pre = True
            self._pre_text = ""

        elif tag == "code":
            if not self._in_pre:
                self._in_code = True

        elif tag == "blockquote":
            self._in_blockquote = True
            # Blockquote content is italic by default (CSS font-style: italic)
            self._italic = True

        elif tag in ("strong", "b"):
            self._bold = True

        elif tag in ("em", "i"):
            self._italic = True

        elif tag == "br":
            run = self._current_para().add_run()
            run._r.append(OxmlElement("w:br"))
            self._last_was_br = True

        elif tag == "img":
            self._embed_image(dict(attrs))

        elif tag == "hr":
            self._paragraph = self.doc.add_paragraph()
            self._paragraph.paragraph_format.space_before = Pt(6)
            self._paragraph.paragraph_format.space_after = Pt(6)
            hr_color = (self._theme.get("color_hr") or "#aaaaaa").lstrip("#").upper()
            hr_sz = str(max(1, round(self._theme.get("size_hr", 0.75) * 8)))
            pPr = self._paragraph._p.get_or_add_pPr()
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = OxmlElement("w:pBdr")
                pPr.append(pBdr)
            bottom = pBdr.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                pBdr.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), hr_sz)
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), hr_color)

        elif tag == "table":
            self._in_table = True
            self._table_rows = []
            self._current_row = []
            self._current_cell_html = ""
            self._in_cell = False
            self._in_th = False
            # Consume any col-widths comment that immediately preceded this table
            self._active_table_col_widths = self._next_table_col_widths
            self._next_table_col_widths = None

        elif tag == "tr":
            self._current_row = []

        elif tag in ("th", "td"):
            self._in_th = tag == "th"
            self._in_cell = True
            self._current_cell_html = ""

    def handle_endtag(self, tag: str) -> None:
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()
        tag = tag.lower()

        # Inside a cell, collect closing inline tags as raw HTML
        if self._in_cell and tag not in ("th", "td"):
            self._current_cell_html += f"</{tag}>"
            return

        if tag in ("h1", "h2", "h3", "h4") and self._paragraph is not None:
            self._apply_section_bar_runs()

        if (
            tag == "h1"
            and self._paragraph is not None
            and self._section_bar_active_tag is None  # section bar replaces the default rule
        ):
            color = self._theme.get("h1_border_color")
            pt = self._theme.get("h1_border_pt", 1.5)
            if color:
                pPr = self._paragraph._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"), "single")
                bot.set(qn("w:sz"), str(max(1, round(pt * 8))))
                bot.set(qn("w:space"), "6")
                bot.set(qn("w:color"), color.lstrip("#").upper())
                pBdr.append(bot)
                pPr.append(pBdr)

        if tag in ("h1", "h2", "h3", "h4", "p"):
            self._section_bar_active_tag = None
            self._paragraph = None

        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
            if self._list_counters:
                self._list_counters.pop()
            self._paragraph = None

        elif tag == "li":
            self._paragraph = None

        elif tag == "dt":
            self._bold = False
            self._paragraph = None

        elif tag == "dd":
            self._paragraph = None

        elif tag == "pre":
            self._in_pre = False
            para = self.doc.add_paragraph(style="Normal")
            pre_size = self._theme.get("font_size_pre", 9.0)
            pre_font = self._theme.get("font_code", "Courier New")
            pre_border_color = self._theme.get("pre_border_color")
            pre_border_pt = self._theme.get("pre_border_pt", 3.0)
            pre_bg = self._theme.get("pre_background_color")
            # Apply background shading (CSS background)
            if pre_bg:
                set_para_shading(para, pre_bg)
            # Apply left accent border (CSS border-left)
            if pre_border_color:
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), str(max(1, round(pre_border_pt * 8))))
                left.set(qn("w:space"), "12")
                left.set(qn("w:color"), pre_border_color.lstrip("#").upper())
                pBdr.append(left)
                pPr.append(pBdr)
                ind = pPr.get_or_add_ind()
                ind.set(qn("w:left"), str(int(10 * 20)))  # 10pt indent
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(10)
            run = para.add_run(self._pre_text.strip())
            run.font.name = pre_font
            run.font.size = Pt(pre_size)
            self._pre_text = ""
            self._paragraph = None

        elif tag == "code":
            if not self._in_pre:
                self._in_code = False

        elif tag == "blockquote":
            self._in_blockquote = False
            self._italic = False
            self._paragraph = None

        elif tag == "div":
            if self._alignment_stack:
                self._alignment_stack.pop()
            self._paragraph = None

        elif tag == "a":
            if self._link_text_buf and self._paragraph is not None:
                if self._current_href:
                    _insert_hyperlink(self._current_para(), self._link_text_buf, self._current_href)
                else:
                    self._write_text(self._current_para(), self._link_text_buf)
            self._current_href = None
            self._link_text_buf = ""

        elif tag in ("strong", "b"):
            self._bold = False

        elif tag in ("em", "i"):
            self._italic = False

        elif tag in ("th", "td"):
            if self._in_table:
                self._current_row.append((self._in_th, self._current_cell_html))
            self._in_cell = False

        elif tag == "tr":
            if self._in_table:
                self._table_rows.append(self._current_row)

        elif tag == "table":
            self._in_table = False
            self._flush_table()

    def _text_width_emu(self) -> int:
        section = self.doc.sections[0]
        return int(section.page_width - section.left_margin - section.right_margin)

    def _embed_image(self, attrs: dict[str, str | None]) -> None:
        """Embed an <img> as a picture: a mermaid:// reference or a file asset."""
        if self._in_cell:
            return
        src = attrs.get("src") or ""
        text_width = self._text_width_emu()

        stream: Any = None
        native_w_emu: int | None = None
        m = _MERMAID_IMG_RE.fullmatch(src)
        if m:
            idx = int(m.group(1))
            if idx >= len(self._mermaid_images):
                return
            png, w_px, _h_px = self._mermaid_images[idx]
            stream = BytesIO(png)
            native_w_emu = w_px * _EMU_PER_PX
        else:
            path = _resolve_asset(src, self._doc_path, self._repo_root)
            if path is None:
                # Unresolved image — fall back to alt text so nothing is silently lost.
                alt = attrs.get("alt")
                if alt:
                    self._write_text(self._current_para(), str(alt))
                logger.warning("docx: could not resolve image %r — skipped.", src)
                return
            stream = str(path)
            try:
                from PIL import Image

                with Image.open(path) as im:
                    native_w_emu = int(im.width * _EMU_PER_PX)
            except Exception:
                native_w_emu = None

        width = text_width if native_w_emu is None else min(native_w_emu, text_width)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run()
        try:
            run.add_picture(stream, width=Emu(width))
        except Exception as exc:
            logger.warning("docx: failed to embed image %r: %s", src, exc)
        self._paragraph = None

    def handle_data(self, data: str) -> None:
        if self._in_pre:
            self._pre_text += data
            return
        if self._in_table:
            self._current_cell_html += data
            return
        self._add_text(data)

    def handle_comment(self, data: str) -> None:
        if self._in_table:
            return
        stripped = data.strip()
        if stripped.lower().startswith("col-widths:"):
            raw = stripped[len("col-widths:") :].strip()
            try:
                widths = [float(v.strip()) for v in raw.split(",") if v.strip()]
                if widths:
                    self._next_table_col_widths = widths
            except ValueError:
                pass

    # ------------------------------------------------------------------
    # Table rendering
    # ------------------------------------------------------------------

    def _flush_table(self) -> None:
        rows = self._table_rows
        if not rows:
            return
        max_cols = max(len(r) for r in rows)
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        try:
            table.style = "Table Normal"
        except KeyError:
            table.style = "Normal Table"

        # Explicitly set all table-level borders to none
        _clear_table_borders(table)

        # Set table to 100% text width with fixed layout and equal column widths.
        # autofit collapses columns in DOTX templates because form-field cells
        # have no content to measure against; fixed layout uses the declared
        # gridCol widths regardless of content.
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        section = self.doc.sections[0]
        text_width_emu = section.page_width - section.left_margin - section.right_margin
        text_width_twips = round(text_width_emu / 914400 * 1440)

        # Build per-column widths. Priority: <!-- col-widths --> comment on this
        # table > table_col_widths config key > equal distribution.
        col_widths_twips: list[int]
        weights = self._active_table_col_widths or self._table_col_widths
        self._active_table_col_widths = None
        if weights and len(weights) == max_cols and sum(weights) > 0:
            total = sum(weights)
            col_widths_twips = [round(text_width_twips * w / total) for w in weights]
            # Correct rounding drift so columns exactly fill the text width
            diff = text_width_twips - sum(col_widths_twips)
            col_widths_twips[-1] += diff
        else:
            col_width_twips = text_width_twips // max_cols
            col_widths_twips = [col_width_twips] * max_cols

        for existing_w in tblPr.findall(qn("w:tblW")):
            tblPr.remove(existing_w)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(text_width_twips))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

        for existing_layout in tblPr.findall(qn("w:tblLayout")):
            tblPr.remove(existing_layout)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        for old in tblPr.findall(qn("w:tblInd")):
            tblPr.remove(old)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "0")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

        # Replace the tblGrid python-docx already created with our own.
        # Must remove first: with fixed layout Word reads tblGrid, and having
        # two of them produces invalid OOXML that corrupts table rendering.
        for old_grid in tbl.findall(qn("w:tblGrid")):
            tbl.remove(old_grid)
        tblGrid = OxmlElement("w:tblGrid")
        for cw in col_widths_twips:
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(cw))
            tblGrid.append(gridCol)
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

        # Cell margins from CSS td { padding }
        tb_pt = self._theme.get("padding_cell_tb_pt", 5.0)
        lr_pt = self._theme.get("padding_cell_lr_pt", 9.0)
        tblCellMar = OxmlElement("w:tblCellMar")
        for side_name, pt_val in (
            ("top", tb_pt),
            ("left", lr_pt),
            ("bottom", tb_pt),
            ("right", lr_pt),
        ):
            mar = OxmlElement(f"w:{side_name}")
            mar.set(qn("w:w"), str(int(pt_val * 20)))  # twips = pt * 20
            mar.set(qn("w:type"), "dxa")
            tblCellMar.append(mar)
        tblPr.append(tblCellMar)

        header_bg = self._theme.get("color_table_header_bg")
        header_text_color = self._theme.get("color_table_header_text")
        header_font_size = self._theme.get("font_size_th")
        body_font_size = self._theme.get("font_size_table")
        row_alt_bg = self._theme.get("color_table_row_alt_bg")
        cell_border_color = self._theme.get("color_table_cell_border", "d5d8dc")
        cell_border_size = self._theme.get("size_table_cell_border", 0.5)
        last_border_color = self._theme.get("color_table_last_row_border", cell_border_color)
        last_border_size = self._theme.get("size_table_last_row_border", 1.0)
        font_body = self._theme.get("font_body")
        font_code = self._theme.get("font_code", "Courier New")
        uppercase_th = self._theme.get("uppercase_th", False)
        letter_spacing_th = self._theme.get("letter_spacing_th_pt")

        n_rows = len(rows)

        for r_idx, row_cells in enumerate(rows):
            is_last_row = r_idx == n_rows - 1
            n_row_cells = len(row_cells)
            for c_idx, (is_header, cell_html) in enumerate(row_cells):
                if c_idx >= max_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                # Explicitly set cell width so fixed-layout tables honour the
                # grid widths rather than falling back to Word's own heuristic.
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW_el = OxmlElement("w:tcW")
                tcW_el.set(qn("w:w"), str(col_widths_twips[c_idx]))
                tcW_el.set(qn("w:type"), "dxa")
                tcPr.append(tcW_el)
                para = cell.paragraphs[0]
                # Zero paragraph spacing so cell padding alone controls whitespace
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                _render_cell_html(
                    para, cell_html.strip(), self._theme, self._write_text, bold_override=is_header
                )

                # Apply alignment to cell paragraphs — Word table cells don't
                # inherit document-level/div alignment the way body text does,
                # so resolve through the same cascade (div > body_text_align).
                if not is_header:
                    word_align = self._effective_alignment()
                    if word_align is not None:
                        para.alignment = word_align

                # Apply body font explicitly (Word table cells don't always inherit Normal)
                if font_body:
                    for run in para.runs:
                        if run.font.name is None or run.font.name not in (font_code, "Courier New"):
                            run.font.name = font_body

                # Apply font sizes
                size = header_font_size if is_header else body_font_size
                if size:
                    for run in para.runs:
                        run.font.size = Pt(size)

                # Header styling
                if is_header:
                    if uppercase_th:
                        for run in para.runs:
                            run.text = run.text.upper()
                    if header_text_color:
                        r, g, b = _hex_to_rgb(header_text_color)
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(r, g, b)
                    if letter_spacing_th:
                        twips = int(letter_spacing_th * 20)
                        for run in para.runs:
                            rPr = run._r.get_or_add_rPr()
                            spacing = OxmlElement("w:spacing")
                            spacing.set(qn("w:val"), str(twips))
                            rPr.append(spacing)
                    if header_bg:
                        set_cell_shading(cell, header_bg)
                else:
                    # Alternating row shading. CSS tr:nth-child(even) counts the
                    # header as child 1, so the shaded body rows are the ones at
                    # odd 0-based table index (r_idx 1, 3, 5, …).
                    if row_alt_bg and r_idx % 2 == 1:
                        set_cell_shading(cell, row_alt_bg)
                    # Bottom border
                    _set_cell_bottom_border(
                        cell,
                        color=last_border_color if is_last_row else cell_border_color,
                        pt=last_border_size if is_last_row else cell_border_size,
                    )

            # Rows shorter than max_cols still have Word cells that need
            # explicit tcW; without it fixed-layout tables render incorrectly.
            for c_idx in range(min(n_row_cells, max_cols), max_cols):
                cell = table.cell(r_idx, c_idx)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW_el = OxmlElement("w:tcW")
                tcW_el.set(qn("w:w"), str(col_widths_twips[c_idx]))
                tcW_el.set(qn("w:type"), "dxa")
                tcPr.append(tcW_el)

        self._paragraph = None


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------


# Named page sizes in mm, portrait (w, h) — matches WeasyPrint/CSS @page sizes.
_PAGE_SIZES_MM = {
    "a3": (297.0, 420.0),
    "a4": (210.0, 297.0),
    "a5": (148.0, 210.0),
    "letter": (215.9, 279.4),
    "legal": (215.9, 355.6),
}
# Default geometry (A4 + PDF theme margins), used when no @page is found.
_DEFAULT_GEOMETRY = {
    "w": 210.0,
    "h": 297.0,
    "top": 25.0,
    "right": 20.0,
    "bottom": 22.0,
    "left": 25.0,
}


def _length_to_mm(token: str) -> float | None:
    """Convert a CSS length (mm/cm/in/pt/px) to mm."""
    m = re.match(r"^([\d.]+)\s*(mm|cm|in|pt|px)?$", token.strip())
    if not m:
        return None
    val = float(m.group(1))
    unit = m.group(2) or "mm"
    return {
        "mm": val,
        "cm": val * 10,
        "in": val * 25.4,
        "pt": val * 25.4 / 72,
        "px": val * 25.4 / 96,
    }[unit]


def _page_geometry(css_text: str | None) -> dict[str, float]:
    """Parse the ``@page { size; margin }`` from theme CSS into mm geometry.

    Mirrors the PDF's page so docx uses the same paper size and margins (and
    therefore the same text width → consistent pagination). Falls back to A4 +
    the standard margins when absent.
    """
    geom = dict(_DEFAULT_GEOMETRY)
    if not css_text:
        return geom
    block = re.search(r"@page\s*\{([^}]*)\}", css_text, re.IGNORECASE)
    if not block:
        return geom
    body = block.group(1)

    size_m = re.search(r"size:\s*([^;]+);", body, re.IGNORECASE)
    if size_m:
        tokens = size_m.group(1).lower().split()
        named = next((t for t in tokens if t in _PAGE_SIZES_MM), None)
        if named:
            w, h = _PAGE_SIZES_MM[named]
            if "landscape" in tokens:
                w, h = h, w
            geom["w"], geom["h"] = w, h
        else:
            lengths = [_length_to_mm(t) for t in tokens]
            lengths = [x for x in lengths if x is not None]
            if len(lengths) >= 2:
                geom["w"], geom["h"] = lengths[0], lengths[1]

    margin_m = re.search(r"margin:\s*([^;]+);", body, re.IGNORECASE)
    if margin_m:
        vals = [_length_to_mm(t) for t in margin_m.group(1).split()]
        vals = [v for v in vals if v is not None]
        if len(vals) == 1:
            geom["top"] = geom["right"] = geom["bottom"] = geom["left"] = vals[0]
        elif len(vals) == 2:
            geom["top"] = geom["bottom"] = vals[0]
            geom["right"] = geom["left"] = vals[1]
        elif len(vals) == 3:
            geom["top"], geom["right"], geom["bottom"] = vals[:3]
            geom["left"] = vals[1]
        elif len(vals) >= 4:
            geom["top"], geom["right"], geom["bottom"], geom["left"] = vals[:4]
    return geom


def _setup_page(doc: Document, geometry: dict[str, float] | None = None) -> None:
    """Set page size and margins to match the PDF layout (from theme @page)."""
    g = geometry or _DEFAULT_GEOMETRY
    section = doc.sections[0]
    section.page_width = Mm(g["w"])
    section.page_height = Mm(g["h"])
    section.top_margin = Mm(g["top"])
    section.right_margin = Mm(g["right"])
    section.bottom_margin = Mm(g["bottom"])
    section.left_margin = Mm(g["left"])


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------


_TWIPS_PER_MM = 1440 / 25.4


def _cfg_mm(value: Any, default_mm: float) -> float:
    """Parse a config length (e.g. ``"10mm"``, ``"1cm"``) to mm."""
    if value is None:
        return default_mm
    mm = _length_to_mm(str(value))
    return mm if mm is not None else default_mm


def _set_para_mark_size(paragraph: Any, half_points: int = 2) -> None:
    """Shrink a paragraph's mark so an empty paragraph takes ~no vertical space."""
    pPr = paragraph._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(half_points))
    rPr.append(sz)


def _tiny_spacer(doc: Document) -> Any:
    """Add a ~zero-height paragraph (separates consecutive floating tables)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _set_para_mark_size(p)
    return p


def _add_floating_table(
    doc: Document,
    *,
    width_twips: int,
    height_mm: float,
    fill: str,
    y_spec: str | None = None,
    y_mm: float | None = None,
    x_mm: float | None = None,
    height_rule: str = "exact",
    inline_indent_twips: int | None = None,
) -> Any:
    """Add a 1×1 borderless shaded table band for the cover page.

    By default the band is floated to an absolute page position (used for the
    top bar/stripe so they sit at the physical page edges the way the PDF
    cover's full-bleed elements do, given ``@page cover { margin: 0 }``).
    With *inline_indent_twips* the band stays in the normal flow instead and
    bleeds into the margins via a negative table indent — used where the PDF
    also lays the element out in flow (e.g. the bottom cover bar).
    """
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Normal"
    except KeyError:
        table.style = "Normal Table"
    _clear_table_borders(table)

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    if inline_indent_twips is None:
        # Absolute page positioning (w:tblpPr) + allow overlap with other floats.
        tblpPr = OxmlElement("w:tblpPr")
        tblpPr.set(qn("w:horzAnchor"), "page")
        tblpPr.set(qn("w:vertAnchor"), "page")
        if x_mm is None:
            tblpPr.set(qn("w:tblpXSpec"), "left")
        else:
            tblpPr.set(qn("w:tblpX"), str(max(1, round(x_mm * _TWIPS_PER_MM))))
        if y_mm is not None:
            tblpPr.set(qn("w:tblpY"), str(max(1, round(y_mm * _TWIPS_PER_MM))))
        else:
            tblpPr.set(qn("w:tblpYSpec"), y_spec or "top")
        overlap = OxmlElement("w:tblOverlap")
        overlap.set(qn("w:val"), "overlap")
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            tblStyle.addnext(tblpPr)
        else:
            tblPr.insert(0, tblpPr)
        tblpPr.addnext(overlap)
    else:
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), str(inline_indent_twips))
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

    # Fixed layout at the exact requested width.
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Zero default cell margins so the shading fills the full band.
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), "0")
        mar.set(qn("w:type"), "dxa")
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)

    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    gridCol = OxmlElement("w:gridCol")
    gridCol.set(qn("w:w"), str(width_twips))
    tblGrid.append(gridCol)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(max(1, int(Mm(height_mm).pt * 20))))
    trHeight.set(qn("w:hRule"), height_rule)
    trPr.append(trHeight)

    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW_el = OxmlElement("w:tcW")
    tcW_el.set(qn("w:w"), str(width_twips))
    tcW_el.set(qn("w:type"), "dxa")
    tcPr.append(tcW_el)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    _set_para_mark_size(p0)

    return table


def _set_cell_margins_mm(cell: Any, top: float, right: float, bottom: float, left: float) -> None:
    """Set explicit tcMar margins (mm) on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, mm_val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), str(max(0, round(mm_val * _TWIPS_PER_MM))))
        mar.set(qn("w:type"), "dxa")
        tcMar.append(mar)
    tcPr.append(tcMar)


def _add_docx_cover_page(
    doc: Document,
    config: dict[str, Any],
    builder: "_DocxBuilder",
    theme: dict[str, Any],
) -> None:
    """Insert a styled cover page for .docx output, then a page break.

    Mirrors the PDF cover design element for element: full-bleed coloured
    bar(s) at the physical page top/bottom, optional accent stripe and logo,
    cover label, title, short divider rule, metadata (author / date) and a
    footer positioned ~14mm from the page bottom — matching the geometry of
    the PDF's ``@page cover { margin: 0 }`` + ``.cover-content`` padding.
    """
    title = config.get("title", "")
    author = config.get("author", "")
    date_str = config.get("date", "")
    label = str(config.get("cover_label", "Report"))
    show_bar = bool(config.get("cover_bar", True))
    bar_pos = str(config.get("cover_bar_position", "top")).lower()
    bar_h = _cfg_mm(config.get("cover_bar_height"), 10.0)
    bar_top_h = _cfg_mm(config.get("cover_bar_top_height"), bar_h)
    bar_bot_h = _cfg_mm(config.get("cover_bar_bottom_height"), bar_h)
    show_stripe = bool(config.get("cover_stripe", False))
    stripe_h = _cfg_mm(config.get("cover_stripe_height"), 120.0)
    stripe_w = _cfg_mm(config.get("cover_stripe_width"), 6.0)
    text_on_bar = bool(config.get("cover_text_on_bar", False))
    show_divider = bool(config.get("cover_divider", True))
    show_footer = bool(config.get("cover_footer", True))
    show_footer_line = bool(config.get("cover_footer_line", True))
    footer_text = config.get("cover_footer_text") or (
        f"{author}  ·  Confidential" if author else ""
    )
    meta_label = str(config.get("cover_meta_label", "Prepared by"))
    meta_author = str(config.get("cover_meta_author", author))

    # Match the PDF cover: bar + title + divider use $primary (color_h1); the
    # label uses the accent (color_h2); meta value is muted grey with a
    # body-coloured bold label.
    primary = (theme.get("color_h1") or theme.get("color_table_header_bg") or "1b4f72").lstrip("#")
    bar_color = primary
    title_color = primary
    accent = (theme.get("color_h2") or theme.get("color_h1") or "2e86c1").lstrip("#")
    label_color = accent
    divider_color = primary  # .cover-divider { border-top: 3pt solid $primary }
    meta_label_color = (theme.get("color_body") or theme.get("color_strong") or "212529").lstrip(
        "#"
    )
    meta_value_color = (theme.get("color_em") or "5d6d7e").lstrip("#")
    font_body = theme.get("font_body")
    text_align = str(config.get("cover_text_align", "left")).lower()
    para_align = WD_ALIGN_PARAGRAPH.RIGHT if text_align == "right" else WD_ALIGN_PARAGRAPH.LEFT

    section = doc.sections[0]
    page_w_mm = section.page_width / 36000
    page_h_mm = section.page_height / 36000
    top_margin_mm = section.top_margin / 36000
    text_width_mm = (section.page_width - section.left_margin - section.right_margin) / 36000
    page_twips = round(section.page_width / 635)

    has_top_bar = show_bar and bar_pos in ("top", "both")
    has_bottom_bar = show_bar and bar_pos in ("bottom", "both")

    # 1. Full-bleed bars + accent stripe, floated to absolute page positions
    #    (the PDF cover has margin: 0, so its bars touch the physical edges).
    if has_top_bar and not text_on_bar:
        _add_floating_table(
            doc, width_twips=page_twips, height_mm=bar_top_h, fill=bar_color, y_spec="top"
        )
        _tiny_spacer(doc)

    if show_stripe:
        stripe_y = bar_top_h if has_top_bar else 10.0
        _add_floating_table(
            doc,
            width_twips=max(1, round(stripe_w * _TWIPS_PER_MM)),
            height_mm=stripe_h,
            fill=accent,
            y_mm=stripe_y,
        )
        _tiny_spacer(doc)

    bar_logo_val = config.get("cover_bar_logo")
    bar_logo_path = (
        _resolve_asset(str(bar_logo_val), builder._doc_path, builder._repo_root)
        if bar_logo_val
        else None
    )

    # 2. Content container. Normally content flows in the body; with
    #    cover_text_on_bar it sits inside the top bar itself (PDF wraps
    #    .cover-content in .cover-bar-wrapper with the primary background).
    container: Any = doc
    first_space_before_pt = 0.0
    base_indent_l_mm, base_indent_r_mm = 3.0, 10.0  # 28mm/30mm padding vs 25mm/20mm margins
    avail_mm = text_width_mm - base_indent_l_mm - base_indent_r_mm
    if text_on_bar and has_top_bar:
        wrap_tbl = _add_floating_table(
            doc,
            width_twips=page_twips,
            height_mm=bar_top_h,
            fill=bar_color,
            y_spec="top",
            height_rule="atLeast",
        )
        wrap_cell = wrap_tbl.rows[0].cells[0]
        # Mirror .cover-content { padding: 50mm 30mm 20mm 28mm; }
        _set_cell_margins_mm(wrap_cell, top=50.0, right=30.0, bottom=20.0, left=28.0)
        container = wrap_cell
        base_indent_l_mm = base_indent_r_mm = 0.0
        avail_mm = page_w_mm - 28.0 - 30.0
        _tiny_spacer(doc)
    else:
        # .cover-content padding-top is 50mm below the top bar; subtract the
        # section top margin since flow content starts there.
        content_top_mm = (bar_top_h if has_top_bar else 0.0) + 50.0
        first_space_before_pt = max(0.0, (content_top_mm - top_margin_mm) * 72.0 / 25.4)

    used_first_cell_para = False

    def _cover_para(space_before: float, space_after: float) -> Any:
        nonlocal used_first_cell_para
        if container is not doc and not used_first_cell_para:
            used_first_cell_para = True
            p = container.paragraphs[0]
        else:
            p = container.add_paragraph()
        p.alignment = para_align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if base_indent_l_mm:
            p.paragraph_format.left_indent = Mm(base_indent_l_mm)
        if base_indent_r_mm:
            p.paragraph_format.right_indent = Mm(base_indent_r_mm)
        return p

    pending_space_before = first_space_before_pt

    # 3. Optional cover logo above the label (PDF: .cover-logo before .cover-label).
    logo_val = config.get("cover_logo")
    logo_path = (
        _resolve_asset(str(logo_val), builder._doc_path, builder._repo_root) if logo_val else None
    )
    if logo_path:
        lp = _cover_para(pending_space_before, 12)
        pending_space_before = 0.0
        run = lp.add_run()
        try:
            logo_w_emu: int | None = None
            try:
                from PIL import Image

                with Image.open(logo_path) as im:
                    logo_w_emu = int(im.width * _EMU_PER_PX)
            except Exception:
                logo_w_emu = None
            max_w_emu = int(Mm(avail_mm))
            if logo_w_emu:
                run.add_picture(str(logo_path), width=Emu(min(logo_w_emu, max_w_emu)))
            else:
                run.add_picture(str(logo_path))
        except Exception as exc:
            logger.warning("docx cover logo embed failed: %s", exc)

    # 4. Cover label (e.g. "REPORT") — small uppercase accent, tracked out.
    #    PDF: .cover-label { 8.5pt / 700 / letter-spacing 2.5pt / 10mm below }.
    if label:
        lp = _cover_para(pending_space_before, 28)  # 10mm ≈ 28pt below
        pending_space_before = 0.0
        run = lp.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(8.5)
        if font_body:
            _apply_font_name(run.font, font_body)
        r, g, b = _hex_to_rgb(label_color)
        run.font.color.rgb = RGBColor(r, g, b)
        # Letter-spacing 2.5pt (val is in twentieths of a point).
        rPr = run._r.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), "50")
        rPr.append(spacing)

    # 5. Title — large bold run in $primary using the body font (NOT the
    #    built-in serif "Title" style, which looks nothing like the PDF).
    #    PDF: .cover-title { 24pt / 700 / 8mm below }.
    title_para = _cover_para(pending_space_before, 23)  # 8mm ≈ 23pt below
    pending_space_before = 0.0
    trun = title_para.add_run(title or "Document")
    trun.bold = True
    trun.font.size = Pt(24)
    if font_body:
        _apply_font_name(trun.font, font_body)
    r, g, b = _hex_to_rgb(title_color)
    trun.font.color.rgb = RGBColor(r, g, b)

    # 6. Short divider rule — 40mm, 3pt, $primary (PDF: .cover-divider).
    #    A bottom-bordered empty paragraph, indented from the far side so the
    #    rule is 40mm wide rather than full-width.
    if show_divider:
        dp = _cover_para(0, 23)  # 8mm below
        _set_para_mark_size(dp)
        gap_mm = max(0.0, avail_mm - 40.0)
        if para_align == WD_ALIGN_PARAGRAPH.RIGHT:
            dp.paragraph_format.left_indent = Mm(base_indent_l_mm + gap_mm)
        else:
            dp.paragraph_format.right_indent = Mm(base_indent_r_mm + gap_mm)
        pPr = dp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "24")  # 3pt
        bot.set(qn("w:space"), "0")
        bot.set(qn("w:color"), divider_color.upper())
        pBdr.append(bot)
        pPr.append(pBdr)

    # 7. Author / date metadata — bold body-coloured label + muted value, no
    #    colons (matches the PDF's "<strong>Prepared by</strong> {author}").
    def _meta_line(bold_label: str, value: str) -> None:
        mp = _cover_para(0, 4)
        lbl = mp.add_run(f"{bold_label} ")
        lbl.bold = True
        lbl.font.size = Pt(10.5)
        if font_body:
            _apply_font_name(lbl.font, font_body)
        lr, lg, lb = _hex_to_rgb(meta_label_color)
        lbl.font.color.rgb = RGBColor(lr, lg, lb)
        val = mp.add_run(value)
        val.font.size = Pt(10.5)
        if font_body:
            _apply_font_name(val.font, font_body)
        vr, vg, vb = _hex_to_rgb(meta_value_color)
        val.font.color.rgb = RGBColor(vr, vg, vb)

    if meta_author:
        _meta_line(meta_label, meta_author)
    if date_str:
        _meta_line("Date", date_str)

    # 8. Bottom bar — laid out in flow after the content, exactly where the
    #    PDF places .cover-bar-bottom (after .cover-content's 20mm bottom
    #    padding). An in-flow full-bleed band also avoids LibreOffice's
    #    clipping of floating tables near the page bottom.
    if has_bottom_bar:
        gap = _tiny_spacer(doc)
        gap.paragraph_format.space_before = Pt(20 * 72 / 25.4)  # 20mm padding-bottom
        left_margin_twips = round(section.left_margin / 635)
        bottom_tbl = _add_floating_table(
            doc,
            width_twips=page_twips,
            height_mm=bar_bot_h,
            fill=bar_color,
            inline_indent_twips=-left_margin_twips,
        )
        if bar_logo_path:
            # The PDF places cover_bar_logo inside the bottom bar, right-aligned
            # with the content edge.
            bp = bottom_tbl.rows[0].cells[0].paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            bp.paragraph_format.right_indent = section.right_margin
            try:
                bp.add_run().add_picture(
                    str(bar_logo_path), height=Mm(max(min(bar_bot_h * 0.7, 8.0), 3.0))
                )
            except Exception as exc:
                logger.warning("docx cover bar logo embed failed: %s", exc)
        _tiny_spacer(doc)

    # 9. Footer (confidentiality notice) — a text frame positioned so its text
    #    sits ~14mm from the physical page bottom, spanning the PDF cover
    #    footer's 28mm→(width−20mm) span, with an optional thin top rule
    #    (PDF: .cover-footer { bottom: 14mm; border-top: 1pt #d5d8dc; 8pt }).
    if show_footer and footer_text:
        fp = doc.add_paragraph()
        fp.alignment = para_align
        pPr = fp._p.get_or_add_pPr()
        framePr = OxmlElement("w:framePr")
        frame_w_mm = max(10.0, page_w_mm - 28.0 - 20.0)
        frame_top_mm = max(0.0, page_h_mm - 14.0 - 7.0)  # ≈ rule + padding + one 8pt line
        framePr.set(qn("w:w"), str(round(frame_w_mm * _TWIPS_PER_MM)))
        framePr.set(qn("w:h"), "0")
        framePr.set(qn("w:hRule"), "auto")
        framePr.set(qn("w:wrap"), "around")
        framePr.set(qn("w:hAnchor"), "page")
        framePr.set(qn("w:x"), str(round(28.0 * _TWIPS_PER_MM)))
        framePr.set(qn("w:vAnchor"), "page")
        framePr.set(qn("w:y"), str(round(frame_top_mm * _TWIPS_PER_MM)))
        pPr.append(framePr)
        if show_footer_line:
            # Thin top rule above the footer text (PDF: 1pt #d5d8dc, 4mm padding).
            pBdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), "8")  # 1pt
            top.set(qn("w:space"), "8")  # ~4pt padding above text
            top.set(qn("w:color"), "D5D8DC")
            pBdr.append(top)
            pPr.append(pBdr)
        run = fp.add_run(footer_text)
        run.font.size = Pt(8)
        if font_body:
            _apply_font_name(run.font, font_body)
        col = config.get("cover_footer_color") or "#7f8c9a"  # PDF .cover-footer colour
        r, g, b = _hex_to_rgb(str(col))
        run.font.color.rgb = RGBColor(r, g, b)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# .dotx content-type patch
# ---------------------------------------------------------------------------


def _patch_compatibility_mode(path: Path) -> None:
    """Upgrade compatibilityMode from 14 (Word 2010) to 15 (Word 2016+).

    python-docx's built-in template ships with compatibilityMode=14, which
    causes Word to open the file in Compatibility Mode.  Patching to 15 tells
    Word the document is fully modern and suppresses the banner.
    """
    tmp = path.with_suffix(".tmp")
    shutil.move(str(path), str(tmp))
    try:
        with zipfile.ZipFile(tmp, "r") as zin:
            with zipfile.ZipFile(path, "w") as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/settings.xml":
                        data = data.replace(
                            b'w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="14"',
                            b'w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="15"',
                        )
                    zout.writestr(item, data)
    except Exception:
        shutil.move(str(tmp), str(path))
        raise
    finally:
        if tmp.exists():
            tmp.unlink()


def _patch_to_dotx(path: Path) -> None:
    """Re-write *path* with the Word Template content type.

    python-docx always saves with the .docx content type. A .dotx differs
    only in one attribute inside ``[Content_Types].xml`` inside the ZIP.
    """
    tmp = path.with_suffix(".tmp")
    shutil.move(str(path), str(tmp))
    try:
        with zipfile.ZipFile(tmp, "r") as zin:
            with zipfile.ZipFile(path, "w") as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        data = data.replace(
                            b"wordprocessingml.document.main+xml",
                            b"wordprocessingml.template.main+xml",
                        )
                    zout.writestr(item, data)
    except Exception:
        shutil.move(str(tmp), str(path))
        raise
    finally:
        if tmp.exists():
            tmp.unlink()


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def _strip_frontmatter(md_content: str) -> str:
    return re.sub(r"^---\s*\n.*?\n---\s*\n", "", md_content, count=1, flags=re.DOTALL)


def _extract_title(md_content: str) -> str | None:
    match = re.search(r"^#\s+(.+)$", md_content, re.MULTILINE)
    if not match:
        return None
    title = match.group(1).strip()
    title = re.sub(r"\*\*(.+?)\*\*", r"\1", title)
    title = re.sub(r"\*(.+?)\*", r"\1", title)
    title = re.sub(r"`(.+?)`", r"\1", title)
    return title


def _strip_leading_h1(md_content: str) -> str:
    return re.sub(r"^#\s+.+\n?", "", md_content, count=1, flags=re.MULTILINE)


def _strip_form_fields_for_docx(md_content: str) -> str:
    """Render PDF ``?[...]`` form markers as a plain fill-in line for Word.

    Word output doesn't support the interactive AcroForm fields the PDF builder
    produces from ``?[...]`` markers, so rather than leaking the literal marker
    text we substitute an underscore fill-in line. ``?[row]``/``?[/row]`` layout
    markers are dropped.
    """
    from .pdf import _FORM_FIELD_RE

    md_content = re.sub(r"^\?\[/?row\]\s*$", "", md_content, flags=re.MULTILINE)
    return _FORM_FIELD_RE.sub("________", md_content)


def _resolve_docx_theme(doc_path: Path | None, repo_root: Path | None) -> dict[str, Any]:
    if doc_path is None or repo_root is None:
        return {}
    result = resolve_docx_theme(doc_path, repo_root)
    return result if result is not None else {}


def _add_page_header_bar(
    doc: Document,
    config: dict[str, Any],
    doc_path: Path | None,
    repo_root: Path | None,
) -> None:
    """Add a coloured header bar with optional text/logos to every page.

    Mirrors the PDF's ``.page-header-bar-fixed``: a full-bleed bar at the
    physical page top (the PDF positions it with negative margin offsets),
    8pt regular text, logos capped at 8mm, content aligned with the page
    margins, and the content area starting ``height + padding`` below the top.
    """
    if not config.get("page_header_bar"):
        return

    color_hex = str(config.get("page_header_bar_color", "#2563eb")).lstrip("#")
    text_color_hex = str(config.get("page_header_bar_text_color", "#ffffff")).lstrip("#")
    height_mm = _cfg_mm(config.get("page_header_bar_height"), 12.0)
    gap_mm = _cfg_mm(config.get("page_header_bar_padding"), 6.0)
    header_text = config.get("header_text", "")
    text_position = str(config.get("header_text_position", "left")).lower()

    single_logo = config.get("page_header_bar_logo") or config.get("header_logo")
    single_logo_position = str(
        config.get("page_header_bar_logo_position") or config.get("header_logo_position", "right")
    ).lower()

    section = doc.sections[0]

    # The PDF bar starts at the physical page top and content begins
    # height + padding below it (@page { margin-top: calc(h + p) }).
    section.header_distance = Mm(0)
    section.top_margin = Mm(height_mm + gap_mm)

    header = section.header
    header.is_linked_to_previous = False
    for para in list(header.paragraphs):
        para._p.getparent().remove(para._p)

    left_margin_twips = round(section.left_margin / 635)
    page_twips = round(section.page_width / 635)

    table = header.add_table(rows=1, cols=3, width=section.page_width)

    # Set the same style as body tables so the style-level tblInd is 0.
    # Without this, Word applies the default table style (tblInd ~108 twips)
    # which shifts the bar right of body content even with explicit tblInd=0.
    try:
        table.style = "Table Normal"
    except KeyError:
        table.style = "Normal Table"

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Full page width, bleeding into both margins via a negative left indent.
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(page_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(-left_margin_twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # Zero default cell margins; the outer cells re-add the page margins so
    # bar content aligns with body content (PDF: padding 0 20mm 0 25mm).
    tblCellMar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        mar = OxmlElement(f"w:{side}")
        mar.set(qn("w:w"), "0")
        mar.set(qn("w:type"), "dxa")
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)

    # 3 slots: left / center / right.
    side_w = round(page_twips * 0.35)
    col_widths = [side_w, page_twips - 2 * side_w, side_w]
    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for cw in col_widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(cw))
        tblGrid.append(gridCol)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    row = table.rows[0]
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Mm(height_mm).pt * 20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    aligns = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)
    for c_idx, cell in enumerate(row.cells):
        set_cell_shading(cell, color_hex)
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        tcW_el = OxmlElement("w:tcW")
        tcW_el.set(qn("w:w"), str(col_widths[c_idx]))
        tcW_el.set(qn("w:type"), "dxa")
        tcPr.append(tcW_el)
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)
        para = cell.paragraphs[0]
        para.alignment = aligns[c_idx]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _set_para_mark_size(para)
    # Align outer-cell content with the page margins.
    _set_cell_margins_mm(row.cells[0], top=0, right=0, bottom=0, left=section.left_margin / 36000)
    _set_cell_margins_mm(row.cells[2], top=0, right=section.right_margin / 36000, bottom=0, left=0)

    slot_paras = {
        "left": row.cells[0].paragraphs[0],
        "center": row.cells[1].paragraphs[0],
        "right": row.cells[2].paragraphs[0],
    }
    logo_height = Mm(max(min(height_mm * 0.7, 8.0), 3.0))  # PDF: .phb-logo max-height 8mm

    if header_text:
        para = slot_paras.get(text_position, slot_paras["left"])
        run = para.add_run(str(header_text))
        run.font.size = Pt(8)  # PDF: .phb-text { font-size: 8pt; } (not bold)
        run.font.color.rgb = RGBColor.from_string(text_color_hex.upper())

    def _slot_picture(pos: str, path: Path) -> None:
        para = slot_paras.get(pos, slot_paras["center"])
        if para.runs:
            para.add_run("  ")
        try:
            para.add_run().add_picture(str(path), height=logo_height)
        except Exception as exc:
            logger.warning("docx header bar logo embed failed: %s", exc)

    if single_logo:
        logo_path = _resolve_asset(str(single_logo), doc_path, repo_root)
        if logo_path:
            _slot_picture(single_logo_position, logo_path)

    for entry in config.get("page_header_bar_logos", []) or []:
        if isinstance(entry, dict):
            lpath = _resolve_asset(str(entry.get("path", "")), doc_path, repo_root)
            pos = str(entry.get("position", "center")).lower()
        else:
            lpath = _resolve_asset(str(entry), doc_path, repo_root)
            pos = "center"
        if lpath:
            _slot_picture(pos, lpath)


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------


_FOOTER_TOKEN_RE = re.compile(r"(\{pages?\})")

_CSS_CONTENT_TOKEN_RE = re.compile(
    r'"([^"]*)"'  # double-quoted string
    r"|'([^']*)'"  # single-quoted string
    r"|counter\(\s*pages\s*\)"
    r"|counter\(\s*page\s*\)"
    r"|string\([^)]*\)"
)


def _css_content_to_text(value: str, date_str: str) -> str | None:
    """Convert a CSS ``content`` value into footer text.

    Quoted strings become literals, ``counter(page)``/``counter(pages)`` become
    the ``{page}``/``{pages}`` tokens the Word footer expands to live fields,
    and ``string(...)`` (the theme's running date) becomes *date_str*.
    """
    if value.strip().lower() in ("none", "normal"):
        return None
    out: list[str] = []
    for m in _CSS_CONTENT_TOKEN_RE.finditer(value):
        if m.group(1) is not None:
            out.append(m.group(1))
        elif m.group(2) is not None:
            out.append(m.group(2))
        elif m.group(0).startswith("string"):
            out.append(date_str)
        elif "pages" in m.group(0):
            out.append("{pages}")
        else:
            out.append("{page}")
    text = "".join(out)
    return text or None


def _css_footer_defaults(css_text: str | None, date_str: str) -> dict[str, tuple[str, float, str]]:
    """Extract the theme's default ``@page`` footer boxes for Word.

    The PDF renders ``@bottom-left/center/right`` margin boxes straight from
    the theme CSS (org name, running date, "Page N of M") even when no
    ``footer_*`` config keys are set. Parse those defaults so the Word footer
    shows the same content. Returns ``{slot: (text, font_size_pt, color_hex)}``.
    """
    if not css_text:
        return {}
    clean = re.sub(r"/\*.*?\*/", "", css_text, flags=re.DOTALL)
    m = re.search(r"@page\s*\{", clean)
    if not m:
        return {}
    # Brace-match the @page block (it contains nested margin-box blocks).
    depth, i = 1, m.end()
    while i < len(clean) and depth:
        if clean[i] == "{":
            depth += 1
        elif clean[i] == "}":
            depth -= 1
        i += 1
    block = clean[m.end() : i - 1]

    result: dict[str, tuple[str, float, str]] = {}
    for slot in ("left", "center", "right"):
        bm = re.search(rf"@bottom-{slot}\s*\{{([^}}]*)\}}", block)
        if not bm:
            continue
        props: dict[str, str] = {}
        for decl in bm.group(1).split(";"):
            if ":" in decl:
                prop, _, val = decl.partition(":")
                props[prop.strip().lower()] = val.strip()
        text = _css_content_to_text(props.get("content", ""), date_str)
        if not text:
            continue
        size_m = re.search(r"([\d.]+)\s*pt", props.get("font-size", ""))
        size = float(size_m.group(1)) if size_m else 7.5
        color_m = re.search(r"#([0-9a-fA-F]{6}|[0-9a-fA-F]{3})\b", props.get("color", ""))
        color = color_m.group(0) if color_m else "#7f8c9a"
        result[slot] = (text, size, color)
    return result


def _add_simple_field(paragraph: Any, instr: str) -> Any:
    """Append a simple Word field (e.g. ``PAGE``/``NUMPAGES``); return its value run."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run2 = paragraph.add_run()
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    run2._r.append(instr_el)

    run3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(sep)

    value_run = paragraph.add_run("1")

    run5 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run5._r.append(end)
    return value_run


def _emit_footer_segment(paragraph: Any, text: str) -> list[Any]:
    """Write *text* into *paragraph*, expanding ``{page}``/``{pages}`` to fields.

    Returns every run created so the caller can apply a consistent footer style.
    """
    runs: list[Any] = []
    for part in _FOOTER_TOKEN_RE.split(text):
        if part == "{page}":
            runs.append(_add_simple_field(paragraph, "PAGE"))
        elif part == "{pages}":
            runs.append(_add_simple_field(paragraph, "NUMPAGES"))
        elif part:
            for i, line in enumerate(part.split("\n")):
                if i > 0:
                    paragraph.add_run().add_break()
                if line:
                    runs.append(paragraph.add_run(line))
    return runs


def _add_footer(
    doc: Document,
    config: dict[str, Any],
    *,
    css_text: str | None = None,
    date_str: str = "",
) -> None:
    """Populate the document footer from footer_left/center/right.

    The three slots are laid out with centre and right tab stops so they mirror
    the PDF ``@bottom-left/center/right`` margin boxes. ``{page}`` and
    ``{pages}`` tokens become live Word page-number fields (the same tokens work
    in the PDF footer). A top border separates the footer from body content.

    Per-slot resolution matches the PDF exactly: an explicitly configured
    ``footer_*`` key wins (6pt, #7f8c9a — the style ``_build_footer_style``
    injects), an empty string suppresses the slot, and an absent key falls back
    to the theme CSS's ``@bottom-*`` default box (its own font-size/colour).
    """
    defaults = _css_footer_defaults(css_text, date_str)

    slots: dict[str, tuple[str, float, str] | None] = {}
    for slot in ("left", "center", "right"):
        cfg_val = config.get(f"footer_{slot}")
        if cfg_val is not None:
            # Explicit config — empty string suppresses the slot (content: none).
            slots[slot] = (str(cfg_val), 6.0, "#7f8c9a") if str(cfg_val) else None
        else:
            slots[slot] = defaults.get(slot)

    if not any(slots.values()):
        return

    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    footer = section.footer

    # Clear the default empty paragraph Word adds
    for para in list(footer.paragraphs):
        p = para._element
        p.getparent().remove(p)

    text_width_emu = int(section.page_width - section.left_margin - section.right_margin)

    para = footer.add_paragraph()

    from docx.enum.text import WD_TAB_ALIGNMENT

    tabs = para.paragraph_format.tab_stops
    tabs.add_tab_stop(Emu(text_width_emu // 2), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Emu(text_width_emu), WD_TAB_ALIGNMENT.RIGHT)

    # Top border (theme @bottom-* boxes: 0.5pt #d5d8dc). The PDF drops the
    # footer border when a page header bar is enabled — mirror that here.
    if not config.get("page_header_bar"):
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), "d5d8dc")
        pBdr.append(top)
        pPr.append(pBdr)

    for i, slot in enumerate(("left", "center", "right")):
        if i > 0:
            para.add_run().add_tab()
        entry = slots[slot]
        if not entry:
            continue
        text, size_pt, color_hex = entry
        r, g, b = _hex_to_rgb(color_hex)
        for run in _emit_footer_segment(para, text):
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(r, g, b)


def _add_plain_header(
    doc: Document,
    config: dict[str, Any],
    doc_path: Path | None,
    repo_root: Path | None,
) -> None:
    """Add header_text / header_logo to the page header when no header bar is used.

    Mirrors the PDF ``@top-left``/``@top-right`` margin boxes so a document that
    only sets ``header_text``/``header_logo`` (without ``page_header_bar``) still
    shows them in Word.
    """
    if config.get("page_header_bar"):
        return  # the coloured bar already renders header text/logo
    header_text = config.get("header_text")
    logo_file = config.get("header_logo")
    if not header_text and not logo_file:
        return

    text_position = str(config.get("header_text_position", "left")).lower()
    logo_position = str(config.get("header_logo_position", "right")).lower()
    logo_path = _resolve_asset(str(logo_file), doc_path, repo_root) if logo_file else None

    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    header = section.header
    for para in list(header.paragraphs):
        para._p.getparent().remove(para._p)

    text_width_emu = int(section.page_width - section.left_margin - section.right_margin)
    para = header.add_paragraph()
    from docx.enum.text import WD_TAB_ALIGNMENT

    tabs = para.paragraph_format.tab_stops
    tabs.add_tab_stop(Emu(text_width_emu // 2), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Emu(text_width_emu), WD_TAB_ALIGNMENT.RIGHT)

    # Hairline under the header content (theme @top-* boxes carry a
    # border-bottom: 0.5pt solid #d5d8dc in the PDF).
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "3")
    bot.set(qn("w:color"), "d5d8dc")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Place text and logo into left/center/right slots via tab stops.
    slots: dict[str, list[Any]] = {"left": [], "center": [], "right": []}

    def _slot_text(pos: str) -> None:
        run = para.add_run(str(header_text))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    def _slot_logo(pos: str) -> None:
        run = para.add_run()
        try:
            run.add_picture(str(logo_path), height=Mm(6))
        except Exception as exc:
            logger.warning("docx header logo embed failed: %s", exc)

    # Build an ordered slot plan, then emit with tabs between left/center/right.
    plan: dict[str, Any] = {}
    if header_text:
        plan[text_position if text_position in slots else "left"] = _slot_text
    if logo_path:
        plan[logo_position if logo_position in slots else "right"] = _slot_logo

    if "left" in plan:
        plan["left"]("left")
    para.add_run().add_tab()
    if "center" in plan:
        plan["center"]("center")
    para.add_run().add_tab()
    if "right" in plan:
        plan["right"]("right")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build(
    rendered_md: str,
    config: dict[str, Any],
    out_path: Path,
    *,
    doc_path: Path | None = None,
    repo_root: Path | None = None,
    output_format: str = "docx",
) -> None:
    """Convert rendered Markdown to a .docx or .dotx file.

    Parameters
    ----------
    rendered_md:
        Jinja2-rendered Markdown string (may include frontmatter).
    config:
        Merged config dict from load_config().
    out_path:
        Destination path for the generated file.
    doc_path:
        Source .md path — used to resolve the CSS theme cascade.
    repo_root:
        Repo root — used to bound the CSS theme cascade.
    output_format:
        ``"docx"`` (default) or ``"dotx"``.  When ``"dotx"``, ``[[field]]``
        markers are converted to Word fields and the saved file is patched to
        the Word Template content type.
    """
    out_path = Path(out_path).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    is_dotx = output_format == "dotx"

    field_type: str | None = None
    cover_page = bool(config.get("cover_page", True))

    if is_dotx:
        ft = str(config.get("dotx_field_type", "form")).lower()
        field_type = ft if ft in ("form", "merge") else "form"

    body = _strip_frontmatter(rendered_md)
    title: str = config.get("title") or _extract_title(body) or out_path.stem
    # Same defaults as the PDF builder so the cover carries identical metadata
    # (the PDF always shows an author and a date).
    author: str = config.get("author", "Document Producer")
    date_str: str = config.get("date") or datetime.date.today().strftime("%-d %B %Y")

    if cover_page:
        # The first H1 becomes the cover title. Without a cover the H1 stays in
        # the body as a Heading 1 — exactly what the PDF does.
        body = _strip_leading_h1(body)

    # Inject the same page breaks the PDF builder uses so both formats break at
    # identical points: APPENDIX section H2s and explicit <!-- pagebreak -->.
    body = _inject_appendix_breaks(body)
    body = _inject_page_breaks(body)
    body = _strip_form_fields_for_docx(body)

    md_engine = markdown.Markdown(extensions=_MD_EXTENSIONS)
    html = md_engine.convert(body)

    # Render mermaid diagrams to embedded PNGs, themed from the same CSS the PDF
    # uses so diagram colours match across formats. Falls back to leaving the
    # code block if cairosvg is unavailable.
    mermaid_theme = None
    css_text: str | None = None
    try:
        from .pdf import _resolve_css
        from ..mermaid import extract_theme_from_css

        css_path = _resolve_css(config, repo_root, doc_path=doc_path)
        if css_path and css_path.exists():
            css_text = css_path.read_text(encoding="utf-8")
            mermaid_theme = extract_theme_from_css(css_text)
    except Exception:
        mermaid_theme = None
    html, mermaid_images = _render_mermaid_to_images(html, mermaid_theme)

    theme = _resolve_docx_theme(doc_path, repo_root)
    doc = Document()
    # Match the PDF's paper size + margins (parsed from the theme's @page) so
    # both formats share the same text width and break at the same points.
    _setup_page(doc, _page_geometry(css_text))

    props = doc.core_properties
    if is_dotx:
        props.title = re.sub(r"\[\[\w+\]\]", "", title).strip()
        if author:
            props.author = re.sub(r"\[\[\w+\]\]", "", author).strip()
    else:
        props.title = title
        if author:
            props.author = author

    _add_page_header_bar(doc, config, doc_path, repo_root)
    _add_plain_header(doc, config, doc_path, repo_root)
    _add_footer(doc, config, css_text=css_text, date_str=date_str)

    if cover_page:
        # The PDF's @page cover rule suppresses every header/footer margin box
        # on the cover — Word's "different first page" does the same. The
        # first-page header/footer parts are created empty (unlinked).
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False

    body_text_align = str(config.get("body_text_align", "")).lower() or None

    raw_col_widths = config.get("table_col_widths")
    table_col_widths: list[float] | None = None
    if isinstance(raw_col_widths, list) and all(
        isinstance(v, (int, float)) for v in raw_col_widths
    ):
        table_col_widths = [float(v) for v in raw_col_widths]

    section_bar: dict[str, Any] | None = None
    if config.get("section_bar"):
        headings_raw = str(config.get("section_bar_headings", "h1,h2"))
        section_bar = {
            "color": str(config.get("section_bar_color", "#2563eb")),
            "text_color": str(config.get("section_bar_text_color", "#ffffff")),
            "text_on_bar": bool(config.get("section_bar_text_on_bar", True)),
            "headings": {h.strip().lower() for h in headings_raw.split(",") if h.strip()},
        }

    builder = _DocxBuilder(
        doc,
        theme=theme,
        field_type=field_type,
        body_text_align=body_text_align,
        table_col_widths=table_col_widths,
        mermaid_images=mermaid_images,
        doc_path=doc_path,
        repo_root=repo_root,
        section_bar=section_bar,
    )

    if cover_page:
        # Both docx and dotx use the richer composable cover; _write_text keeps
        # [[field]] markers working in dotx output.
        _add_docx_cover_page(
            doc,
            {**config, "title": title, "author": author, "date": date_str},
            builder,
            theme,
        )

    # Content after the cover starts the pagination baseline — the first body
    # heading must not force an extra page break on top of the cover's.
    builder.mark_content_start()
    builder.feed(html)

    doc.save(str(out_path))

    # Patch the document theme XML to match the CSS font — ensures LibreOffice
    # and other renderers show the correct font family instead of Calibri/Cambria.
    font_body = theme.get("font_body")
    if font_body:
        patch_docx_theme_fonts(out_path, font_body)

    _patch_compatibility_mode(out_path)

    if is_dotx:
        _patch_to_dotx(out_path)
